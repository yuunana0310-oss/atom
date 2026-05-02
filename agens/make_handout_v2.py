# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ===== カラーパレット（白黒印刷対応）=====
C_NAVY    = RGBColor(0x1A, 0x1A, 0x1A)   # 黒（メイン見出し）
C_TEAL    = RGBColor(0x55, 0x55, 0x55)   # 濃グレー（サブ見出し）
C_SKY     = RGBColor(0xF0, 0xF0, 0xF0)   # 薄グレー（背景1）
C_MINT    = RGBColor(0xE8, 0xE8, 0xE8)   # やや濃グレー（背景2）
C_ORANGE  = RGBColor(0x33, 0x33, 0x33)   # 警告ボックス（黒）
C_RED     = RGBColor(0x00, 0x00, 0x00)   # 重要（黒）
C_GREEN   = RGBColor(0x33, 0x33, 0x33)   # チェック（黒）
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY1   = RGBColor(0x44, 0x44, 0x44)
C_GRAY2   = RGBColor(0xF5, 0xF5, 0xF5)
C_NAVY_L  = RGBColor(0x55, 0x55, 0x55)   # 中グレー

def rgb_hex(color):
    return '%02X%02X%02X' % (color[0], color[1], color[2])

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_border_bottom(para, color_hex='007A87', sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_shading(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    pPr.append(shd)

def set_table_no_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_border(table, color_hex='007A87', sz=8):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color_hex)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_section_heading(doc, text, icon=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_para_shading(p, C_NAVY)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f'  {icon}  {text}' if icon else f'  {text}')
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = C_WHITE
    return p

def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_para_shading(p, C_TEAL)
    run = p.add_run(f'  {text}')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_WHITE
    return p

def add_note_box(doc, text, box_type='warn'):
    # warn=オレンジ, info=ティール, ok=緑
    if box_type == 'warn':
        bg, icon = C_ORANGE, '⚠'
    elif box_type == 'info':
        bg, icon = C_TEAL, '★'
    else:
        bg, icon = C_GREEN, '✓'
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    set_para_shading(p, bg)
    run = p.add_run(f'  {icon}  {text}')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_WHITE
    return p

def add_bullet(doc, text, size=10, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)
    run_bullet = p.add_run('▸ ')
    run_bullet.font.color.rgb = C_TEAL
    run_bullet.font.size = Pt(size)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_step(doc, step_label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(3)
    # STEPバッジ風
    run1 = p.add_run(f' {step_label} ')
    run1.bold = True
    run1.font.size = Pt(9)
    run1.font.color.rgb = C_WHITE
    # 背景をティールにするため段落シェーディングは使えないので色付きrunで代用
    # 実用上はrun背景色（highlight）は限定色のみ対応のため、太字カラーで表現
    run1.font.color.rgb = C_TEAL
    run1.font.size = Pt(10)
    run1.bold = True
    run2 = p.add_run(f'  {text}')
    run2.font.size = Pt(10)
    return p

# ============================================================
doc = Document()

section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# ===== ヘッダーバー（タイトルブロック）=====
# タイトル背景段落
p_bg = doc.add_paragraph()
set_para_shading(p_bg, C_NAVY)
p_bg.paragraph_format.space_before = Pt(0)
p_bg.paragraph_format.space_after = Pt(0)
p_bg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_bg.add_run('  新人職員研修　配布資料  ')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = C_WHITE

p_sub = doc.add_paragraph()
set_para_shading(p_sub, C_TEAL)
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(2)
r = p_sub.add_run('移乗・移動介助の基本　〜安全・安楽・自立支援を実践する〜')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = C_WHITE

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_info.paragraph_format.space_after = Pt(8)
r = p_info.add_run('対象：看護職員・介護職員　　所要時間：約30分　　医療法人 研修・教育委員会')
r.font.size = Pt(9)
r.font.color.rgb = C_GRAY1

# ===== 1. 研修目標 =====
add_section_heading(doc, '本日の研修目標', '🎯')

goals = [
    '移乗・移動介助の基本原則とボディメカニクスを理解する',
    '車椅子・ベッド間の移乗手順を正しく実施できる',
    '歩行介助・立ち上がり介助を安全に行える',
    '転倒・転落リスクを予測し予防策を講じられる',
    '患者の尊厳と残存能力を尊重した介助ができる',
]
nums = ['①','②','③','④','⑤']
for i, g in enumerate(goals):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'  {nums[i]} ')
    r1.bold = True
    r1.font.color.rgb = C_NAVY
    r1.font.size = Pt(10.5)
    r2 = p.add_run(g)
    r2.font.size = Pt(10.5)
# 上を書き直す
# 既に追加済みなので消してやり直す
# → 一旦このまま継続

add_note_box(doc, '患者の「安全・安楽・自立支援」が介助の3原則です', 'info')

# ===== 2. 介助の3原則 =====
add_section_heading(doc, '介助の3原則', '📋')

table = doc.add_table(rows=2, cols=3)
set_table_border(table, '007A87', 6)
headers = [('安 全', C_NAVY), ('安 楽', C_TEAL), ('自立支援', RGBColor(0x0B, 0x6E, 0x37))]
for i, (h, c) in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = C_WHITE

contents = [
    '・転倒・転落・皮膚損傷を防ぐ\n・環境整備を必ず行う\n・2人介助の基準を遵守する',
    '・患者の苦痛・不安を最小化\n・声かけ・同意を必ず得る\n・プライバシーを守る',
    '・できることは自分で行ってもらう\n・残存能力を活かす介助\n・ADL維持・向上を目指す',
]
bg_colors = [C_SKY, C_MINT, RGBColor(0xF0, 0xF8, 0xF0)]
for i, (text, bg) in enumerate(zip(contents, bg_colors)):
    cell = table.rows[1].cells[i]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

add_note_box(doc, '声かけ・同意・プライバシー保護は必ず実施！', 'warn')

# ===== 3. ボディメカニクス =====
add_section_heading(doc, 'ボディメカニクスの基本（腰痛予防・安全介助）', '🔧')

bm = [
    ('① 支持基底面を広げる', '足を肩幅に開き、安定した姿勢を確保する'),
    ('② 重心を低くする',     '膝を曲げて腰を落とし、重心を下げる'),
    ('③ 重心を近づける',     '患者との距離を縮め、密着して介助する'),
    ('④ 大きな筋群を使う',   '腰だけでなく脚・体幹の大きな筋肉を使う'),
    ('⑤ 身体をひとつにまとめる', '患者の手足を体に引き付けてから動かす'),
    ('⑥ ねじり動作を避ける', '足を向けてから体を回旋させる'),
]
# 2列テーブルで見やすく
bm_table = doc.add_table(rows=3, cols=2)
set_table_border(bm_table, 'DDDDDD', 4)
for idx, (label, desc) in enumerate(bm):
    row_i = idx // 2
    col_i = idx % 2
    cell = bm_table.rows[row_i].cells[col_i]
    set_cell_bg(cell, C_SKY if row_i % 2 == 0 else C_WHITE)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + '\n')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = C_NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(9.5)

# ===== 4. 介助前の確認事項 =====
add_section_heading(doc, '介助前の確認事項（リスクアセスメント・環境整備）', '✅')

conf_table = doc.add_table(rows=2, cols=2)
set_table_border(conf_table, '007A87', 6)
for i, (h, c) in enumerate([('患者確認', C_NAVY), ('環境整備', C_TEAL)]):
    cell = conf_table.rows[0].cells[i]
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_WHITE

patient_items = '・患者の理解・同意を得る\n・バイタルサイン・体調確認\n・点滴・ドレーン・カテーテル確認\n・麻痺・疼痛・可動域の確認\n・体重・介助レベル確認'
env_items = '・ベッドの高さ調整（介助しやすい高さ）\n・ブレーキ・ストッパーの確認\n・床の濡れ・障害物の除去\n・スリッパ→滑り止め靴に変更\n・プライバシーカーテンを閉める'
for i, (text, bg) in enumerate(zip([patient_items, env_items], [C_SKY, C_MINT])):
    cell = conf_table.rows[1].cells[i]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

add_note_box(doc, '2人介助基準：全介助・体重60kg以上・不安定な患者は必ず2名で対応', 'warn')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('  基本手順：  ')
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = C_NAVY
# フロー矢印
flow = ['確認', '声かけ', '環境整備', '介助', '観察', '記録']
for j, step in enumerate(flow):
    rb = p.add_run(step)
    rb.bold = True
    rb.font.color.rgb = C_TEAL
    rb.font.size = Pt(10)
    if j < len(flow) - 1:
        ra = p.add_run(' → ')
        ra.font.color.rgb = C_GRAY1
        ra.font.size = Pt(10)

# ===== 5. 立ち上がり介助 =====
add_section_heading(doc, '立ち上がり介助（ベッド端座位 → 立位）', '🧍')

steps = [
    ('STEP 1', '端座位を確保する', 'ベッド端に座り、足底を床につける（足が浮く場合はフットレスト使用）'),
    ('STEP 2', '前傾姿勢をとる',   '体を前に傾け重心を前方に移動。介助者は腰部を支持する'),
    ('STEP 3', '立ち上がりを誘導', '「せーの」の声かけで一緒に立ち上がる。脇・腰を支えながら立位へ'),
    ('STEP 4', '立位の安定確認',   'めまい・ふらつきがないか確認。必要なら手すり・歩行器を使用'),
]
step_table = doc.add_table(rows=4, cols=2)
set_table_no_border(step_table)
step_table.columns[0].width = Cm(2.5)
for idx, (step, label, desc) in enumerate(steps):
    cell_l = step_table.rows[idx].cells[0]
    cell_r = step_table.rows[idx].cells[1]
    set_cell_bg(cell_l, C_NAVY if idx % 2 == 0 else C_TEAL)
    set_cell_bg(cell_r, C_SKY if idx % 2 == 0 else C_MINT)
    pl = cell_l.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl1 = pl.add_run(step + '\n')
    rl1.bold = True
    rl1.font.size = Pt(8.5)
    rl1.font.color.rgb = C_WHITE
    rl2 = pl.add_run(label)
    rl2.bold = True
    rl2.font.size = Pt(8.5)
    rl2.font.color.rgb = C_WHITE
    pr = cell_r.paragraphs[0]
    rr = pr.add_run(desc)
    rr.font.size = Pt(10)

# ===== 6. 移乗手順 =====
add_section_heading(doc, 'ベッド ⇔ 車椅子 移乗手順（片麻痺患者）', '♿')

trans_table = doc.add_table(rows=2, cols=2)
set_table_border(trans_table, '555555', 6)
for i, (h, c) in enumerate([('ベッド → 車椅子', C_NAVY), ('車椅子 → ベッド', C_TEAL)]):
    cell = trans_table.rows[0].cells[i]
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_WHITE

bw_items = '1. 車椅子を健側に約30°の角度で設置\n2. ブレーキ確認・フットレストを上げる\n3. 患者を端座位にし足底を床につける\n4. 前傾姿勢→健側の手で介助しながら立上がり\n5. ゆっくり方向転換して車椅子に着座\n6. フットレストを下げ姿勢を整える'
wb_items = '（逆の手順で実施）\n1. 健側に30°でベッドに対して配置\n2. ブレーキ・フットレスト確認\n3. 前傾姿勢から立ち上がりを誘導\n4. 方向転換（介助者は患側に位置）\n5. ベッドに手をつかせてゆっくり着座\n6. 体位を整え座位バランスを確認'
for i, (text, bg) in enumerate(zip([bw_items, wb_items], [C_SKY, C_MINT])):
    cell = trans_table.rows[1].cells[i]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

add_note_box(doc, '患者の「できること」を最大限活かす介助を心がけましょう', 'ok')

# ===== 7. 歩行介助 =====
add_section_heading(doc, '歩行介助の基本', '🚶')

walk_table = doc.add_table(rows=2, cols=4)
set_table_border(walk_table, '555555', 4)
walk_headers = [('開始前', C_NAVY), ('出発時', C_TEAL), ('歩行中', C_NAVY), ('着座時', C_TEAL)]
for i, (h, c) in enumerate(walk_headers):
    cell = walk_table.rows[0].cells[i]
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_WHITE

walk_contents = [
    '・靴・履物確認\n・バイタル確認\n・補助具の用意',
    '・立位保持で安定確認\n・めまい確認\n・ゆっくり歩き始める',
    '・声かけしながら一定ペース\n・障害物に注意\n・急な向き変え禁止',
    '・椅子/ベッドに触れてから\n・ゆっくり膝を曲げて\n・着座確認',
]
for i, (text, bg) in enumerate(zip(walk_contents, [C_SKY, C_MINT, C_SKY, C_MINT])):
    cell = walk_table.rows[1].cells[i]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('  介助者の位置：患側（麻痺のある側）の斜め後方　片手は腰または移乗ベルトを握る')
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = C_NAVY

# ===== 8. 全介助（2名） =====
add_section_heading(doc, '全介助の移乗（2名介助）', '👥')

two_table = doc.add_table(rows=2, cols=2)
set_table_border(two_table, '555555', 6)
for i, (h, c) in enumerate([('リーダー役（頭側）', C_NAVY), ('サポート役（足側）', C_TEAL)]):
    cell = two_table.rows[0].cells[i]
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_WHITE

leader = '・上半身・頭頸部を支持\n・声かけ・タイミングの指示を出す\n・「いちにのさん」で全員同時に動く\n・気道・カテーテル類の安全を確認'
support = '・下肢・臀部を支持\n・リーダーの指示に従う\n・足先から骨盤まで確実に保持\n・移動方向を確認してから動く'
for i, (text, bg) in enumerate(zip([leader, support], [C_SKY, C_MINT])):
    cell = two_table.rows[1].cells[i]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

add_note_box(doc, '適応：全介助患者 ／ 体重60kg以上 ／ 不安定な患者 ／ 初回移乗', 'warn')

# ===== 9. 転倒・転落予防 =====
add_section_heading(doc, '転倒・転落予防', '⚠')

fall_table = doc.add_table(rows=2, cols=3)
set_table_border(fall_table, '555555', 6)
for i, (h, c) in enumerate([('患者側リスク', C_NAVY), ('環境側リスク', C_TEAL), ('予防策', C_NAVY)]):
    cell = fall_table.rows[0].cells[i]
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_WHITE

r1 = '・筋力低下・バランス障害\n・認知症・見当識障害\n・薬剤影響（睡眠薬等）\n・起立性低血圧\n・排泄への切迫感'
r2 = '・床の濡れ・障害物\n・不適切な照明\n・ベッド高さ不適切\n・ナースコール未設置\n・滑りやすい履物'
r3 = '・転倒リスクスコア評価（毎日）\n・離床センサー設置\n・低床ベッドの使用\n・ナースコールを手の届く位置へ\n・排泄パターン把握しラウンド強化'
bgs = [C_SKY, C_MINT, C_GRAY2]
for i, (text, bg) in enumerate(zip([r1, r2, r3], bgs)):
    cell = fall_table.rows[1].cells[i]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9)

add_note_box(doc, '転倒発生時：患者の安全確保 → ナースコール → 状態確認 → 報告 → 記録', 'warn')

# ===== 10. よくあるミス =====
add_section_heading(doc, 'よくあるミスと対策', '🔍')

err_table = doc.add_table(rows=1, cols=3)
set_table_border(err_table, 'AAAAAA', 4)
for i, (h, c) in enumerate([('よくあるミス', C_NAVY), ('起こりうる問題', C_TEAL), ('対 策', C_NAVY)]):
    cell = err_table.rows[0].cells[i]
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_WHITE

mistakes = [
    ('× 声かけを忘れる', '患者の心理的不安・筋緊張が増す', '○ 必ず「○○しますね」と伝えてから開始'),
    ('× ブレーキを忘れる', '移乗中に車椅子が動いて転落', '○ 移乗前の確認を手順書通りに実施'),
    ('× 腰だけで持ち上げる', '介助者の腰痛・ぎっくり腰', '○ ボディメカニクスを意識。足・体幹を使う'),
    ('× 急ぎすぎる', '患者が不安・バランス崩す', '○ 患者のペースに合わせ、ゆっくり確実に'),
    ('× 1人で無理をする', '患者・介助者ともに受傷リスク', '○ 迷ったら必ず同僚に声をかける'),
    ('× カテーテルを忘れる', 'ドレーン抜去・皮膚損傷', '○ 移乗前にルート類の長さ・固定を確認'),
]
for idx, (m1, m2, m3) in enumerate(mistakes):
    row = err_table.add_row()
    bg = C_GRAY2 if idx % 2 == 0 else C_WHITE
    for j, (text, extra_color) in enumerate([(m1, C_NAVY), (m2, C_GRAY1), (m3, C_NAVY)]):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = extra_color
        run.bold = (j == 0 or j == 2)

# ===== 11. 確認テスト =====
add_section_heading(doc, '確認テスト', '📝')

qs = [
    ('Q1', 'ボディメカニクスの目的は何ですか？',
     '①患者の安楽　②介助者の腰痛予防　③両方　④どちらでもない'),
    ('Q2', '片麻痺患者への車椅子移乗で、車椅子を置く位置はどちら側？',
     '①患側（麻痺側）　②健側（麻痺のない側）　③どちらでも同じ'),
    ('Q3', '2名介助が必要な条件を1つ答えてください。',
     '（　　　　　　　　　　　　　　　　　　　　　　　　　　　　）'),
    ('Q4', '移乗介助前に必ず確認すべき3つのことを答えてください。',
     '（１）　　　　　　　　　（２）　　　　　　　　　（３）　　　　　　　　'),
    ('Q5', '転倒が発生した際の対応の優先順位を答えてください。',
     '（　　　　　　　　　　　　　　　　　　　　　　　　　　　　）'),
]
for q_num, q_text, options in qs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f'  {q_num}　')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = C_NAVY
    r2 = p.add_run(q_text)
    r2.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.5)
    p2.paragraph_format.space_after = Pt(5)
    r3 = p2.add_run(options)
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = C_GRAY1

# ===== 12. まとめ =====
add_section_heading(doc, '本日のまとめ', '📌')

summary = [
    ('1', '介助の3原則',   '安全・安楽・自立支援を常に意識する'),
    ('2', 'ボディメカニクス', '正しい姿勢で患者も介助者も守る'),
    ('3', '確認・声かけ',  '介助前の確認と患者への声かけを徹底する'),
    ('4', '移乗手順の遵守', '車椅子の角度・ブレーキ・フットレストを確認'),
    ('5', '2名介助の判断', '一人で無理せず、迷ったら必ず声をかける'),
    ('6', '記録・報告',    '気づいたことはすぐに記録・報告・連絡'),
]
sum_table = doc.add_table(rows=6, cols=2)
set_table_no_border(sum_table)
sum_table.columns[0].width = Cm(4.5)
for idx, (num, key, val) in enumerate(summary):
    cl = sum_table.rows[idx].cells[0]
    cr = sum_table.rows[idx].cells[1]
    bg = C_SKY if idx % 2 == 0 else C_MINT
    set_cell_bg(cl, C_NAVY if idx % 2 == 0 else C_TEAL)
    set_cell_bg(cr, bg)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl1 = pl.add_run(num + '　')
    rl1.bold = True
    rl1.font.size = Pt(11)
    rl1.font.color.rgb = RGBColor(0xFF, 0xE0, 0x80)
    rl2 = pl.add_run(key)
    rl2.bold = True
    rl2.font.size = Pt(10)
    rl2.font.color.rgb = C_WHITE
    pr = cr.paragraphs[0]
    rr = pr.add_run(val)
    rr.font.size = Pt(10)

# フィニッシュメッセージ
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fin.paragraph_format.space_before = Pt(10)
set_para_shading(p_fin, C_TEAL)
r_fin = p_fin.add_run('患者の「安全・安楽・自立支援」のために、今日から実践しましょう！')
r_fin.bold = True
r_fin.font.size = Pt(12)
r_fin.font.color.rgb = C_WHITE

doc.save('e:/新人/移乗移動介助_配布資料.docx')
print('保存完了: e:/新人/移乗移動介助_配布資料.docx')
