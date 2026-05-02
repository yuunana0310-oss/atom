from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== ユーティリティ =====

def set_run_font(run, font='Meiryo'):
    run.font.name = font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font)
    rPr.insert(0, rFonts)

def set_shading(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def set_border_left(para, color='000000', sz=24):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def set_border_bottom(para, color='000000', sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)

def p_text(doc, text, size=11, bold=False, color='111111',
           indent=0, sp_before=2, sp_after=3, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after = Pt(sp_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)
    return p

def slide_header(doc, num, title, color_hex):
    """スライド番号 + タイトルのセクション区切り"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    set_shading(p, color_hex)
    r1 = p.add_run(f'  スライド {num}　')
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_run_font(r1)
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_run_font(r2)

def note(doc, text):
    """グレー背景の注釈（ページ参照など）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    set_shading(p, 'F0F0F0')
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_run_font(run)

def script(doc, text, indent=0):
    """発表原稿本文"""
    p = p_text(doc, text, size=11, indent=indent, sp_before=2, sp_after=4)
    return p

def emphasis(doc, text, indent=0):
    """強調読み（太字）"""
    p = p_text(doc, text, size=11, bold=True, color='111111',
               indent=indent, sp_before=2, sp_after=4)
    return p

def spacer(doc, size=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('')
    run.font.size = Pt(size)
    set_run_font(run)


# ===== ドキュメント生成 =====

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

# ===== 表紙 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('発表原稿（スライド用）')
run.bold = True; run.font.size = Pt(20); set_run_font(run)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('病院スタッフ必修研修　個人情報保護・ハラスメント予防・法令遵守')
run2.font.size = Pt(12); set_run_font(run2)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('目安：約20分（早読み）　全13スライド対応版')
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
set_run_font(run3)
set_border_bottom(p3)

# 凡例
spacer(doc, 6)
p_leg = doc.add_paragraph()
p_leg.paragraph_format.left_indent = Cm(0.3)
set_shading(p_leg, 'F5F5F5')
run_l = p_leg.add_run(
    '【読み方】　★ = やや強調してゆっくり読む　／ = 少し間を置く\n'
    '　　　　　　太字 = 特に丁寧に読む　｜　→ インデント行 = 箇条書きを読む行'
)
run_l.font.size = Pt(9)
run_l.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_run_font(run_l)


# ============================================================
# スライド 1: 表紙
# ============================================================
slide_header(doc, 1, '表紙', '1A1A2E')

note(doc, '※ 開始前に表示しておく。参加者が着席したら話し始める')

script(doc, 'みなさん、お疲れのところお集まりいただきありがとうございます。')
script(doc, '本日は約30分、研修の時間をいただきます。')
script(doc, 'テーマはスライドにあります通り、「個人情報保護」「ハラスメント予防」「法令遵守」の3つです。')
script(doc, 'どれも「知っているつもり」になりやすい内容ですが、実際の現場で問題になるのは「知らなかった」ではなく「まあいいか」と思った瞬間です。／')
script(doc, '今日はそのポイントをしっかり確認していきましょう。では始めます。')


# ============================================================
# スライド 2: アジェンダ
# ============================================================
slide_header(doc, 2, 'アジェンダ', '1A1A2E')

note(doc, '※ スライドの4ブロックを指しながら説明')

script(doc, '本日の流れをご確認ください。')
script(doc, 'まずPART 01で個人情報保護、次にPART 02でハラスメント予防——こちらは今話題のカスハラを重点的に扱います。')
script(doc, 'PART 03で法令遵守をお伝えし、最後にセルフチェックリストで締めくくります。')
emphasis(doc, '★ 今日のゴールは3つだけです。')
script(doc, '患者情報を「なぜ・どう守るか」が言えること。／カスハラの定義と対応を知っていること。／違反のリスクを理解して、迷ったとき相談できること。この3点です。', indent=0.5)
script(doc, '配布資料も合わせてご活用ください。')


# ============================================================
# スライド 3: 個人情報 — 定義
# ============================================================
slide_header(doc, 3, '個人情報保護　―　定義', '1B4F8A')

script(doc, 'まず「個人情報」とは何かを整理します。')
script(doc, '左側をご覧ください。氏名・生年月日・住所・電話番号、これが一般的な個人情報です。')
script(doc, '右側——病院で扱うのは「要配慮個人情報」と呼ばれる、さらに厳格な情報です。')
emphasis(doc, '★ 診断名、病歴、精神疾患の有無、感染症の罹患——これらは漏れただけで患者さんの人生に深刻な影響を与えかねない情報です。')
script(doc, 'だから病院では一般企業より格段に厳しい情報管理が求められています。')


# ============================================================
# スライド 4: 個人情報 — 5原則
# ============================================================
slide_header(doc, 4, '個人情報保護　―　5原則', '1B4F8A')

note(doc, '※ 番号を指しながら順番に読む')

script(doc, '患者情報の取り扱いには5つの原則があります。')
for num, pr, detail in [
    ('1', '目的外使用禁止',    '診療目的のみに使用する。それ以外の目的で使ってはいけません。'),
    ('2', '最小限のアクセス',  '自分の担当患者以外のカルテを開くだけで違反になります。'),
    ('3', '第三者提供の制限',  '家族からの問い合わせでも、本人同意なく情報を渡すことは原則できません。'),
    ('4', '安全管理措置',      '書類の放置、PCのロック忘れ——これも違反になり得ます。'),
    ('5', '漏洩時の即時報告',  '問題が起きたとき、隠すより早く報告するほうが被害を最小化できます。'),
]:
    script(doc, f'原則{num}、{pr}。{detail}', indent=0.5)

emphasis(doc, '★ 特に2番——「担当外のカルテをちょっと見る」だけでアウトです。意識してください。')


# ============================================================
# スライド 5: 個人情報 — NG10
# ============================================================
slide_header(doc, 5, '個人情報保護　―　NG10', '1B4F8A')

note(doc, '※ 全部は読まない。赤背景（★）の3項目を重点的に話す')

script(doc, 'やってはいけないことをまとめました。')
emphasis(doc, '★ 特に重要な3つだけ確認します。')
script(doc, '1つ目——私用スマホでの撮影・保存。「メモ代わりに」でも完全にアウトです。撮影した瞬間に違反になります。', indent=0.5)
script(doc, '2つ目——SNSへの投稿。名前を書いていなくても、診断名や状況から特定できれば問題になります。', indent=0.5)
script(doc, '3つ目——退職後の情報利用。辞めたあとも守秘義務は続きます。これは一生です。', indent=0.5)
script(doc, 'その他、廊下での会話、施錠忘れ、誤送信——すべて違反の入口です。スライドで確認しておいてください。')
emphasis(doc, '★ 違反した場合は、1年以下の懲役または50万円以下の罰金という刑事罰があります。')


# ============================================================
# スライド 6: ハラスメント種類
# ============================================================
slide_header(doc, 6, 'ハラスメント予防　―　種類', 'E67E22')

note(doc, '※ 表を上から順に指しながら説明。カスハラで少し速度を落とす')

script(doc, '次にハラスメント予防です。まず種類を整理します。')
script(doc, 'パワハラは優越的な立場を使った言動です。「なんでこんなこともできないんだ」という怒鳴り声——指導ではなくパワハラです。', indent=0.5)
script(doc, 'セクハラは性的な言動。外見や体型への不用意な一言も含まれます。', indent=0.5)
script(doc, 'マタハラは妊娠・育休に関わるもの。「このタイミングで妊娠するの？」は一言でアウトです。', indent=0.5)
emphasis(doc, '★ そして今日の重点、カスハラです。患者・家族から受けるハラスメント——次のスライドで詳しく見ます。')


# ============================================================
# スライド 7: カスハラ — 定義・具体例
# ============================================================
slide_header(doc, 7, 'カスハラ　―　定義・具体例', 'C0392B')

note(doc, '※ 定義ボックスをゆっくり読み上げる。その後、具体例は4ブロックを軽く流す')

script(doc, '厚生労働省の定義を読みます。')
emphasis(doc,
    '★「顧客・患者・その家族等が行う、社会通念上不相当な言動で、'
    '労働者の就業環境が害されるもの」——これが2024年に示された定義です。')
script(doc, '病院は特にカスハラが起きやすい職場です。患者さんやご家族は不安や不満を抱えて来院します。その感情が時にスタッフへの不当な言動に向かうことがあります。')
script(doc, '具体例を見てください。')
for t, d in [
    ('暴言・脅迫系',  '「訴えるぞ」「SNSで晒す」という脅し、大声の怒鳴り。'),
    ('不当要求系',   '土下座の強要、医学的に不必要な処置の要求。'),
    ('身体的行為系', '物を投げる、叩く——これは最悪の場合、刑事事件になります。'),
    ('長時間拘束系', '何時間も居座ってクレームを繰り返す。これも立派なカスハラです。'),
]:
    script(doc, f'{t}——{d}', indent=0.5)


# ============================================================
# スライド 8: カスハラ — 正当クレームとの違い
# ============================================================
slide_header(doc, 8, 'カスハラ　―　正当クレームとの違い', 'C0392B')

note(doc, '※ 左（グリーン）と右（レッド）を交互に指して対比させる')

emphasis(doc, '★ ここが一番大切なポイントです。正当なクレームとカスハラは違います。')
script(doc, '左側、グリーンの欄——待ち時間への不満、診療方針への疑問、改善要望——これらは正当なクレームです。丁寧に受け止めて、真摯に対応しなければなりません。')
script(doc, '右側、レッドの欄——怒鳴る、脅す、土下座を強要する、暴力をふるう——これはカスハラです。')
emphasis(doc, '★ 「患者さんだから」と我慢し続ける必要はありません。受け入れないのが正しい姿勢です。')
script(doc, 'この区別ができることが、適切な対応の第一歩です。')


# ============================================================
# スライド 9: カスハラ — 対応フロー
# ============================================================
slide_header(doc, 9, 'カスハラ　―　対応フロー', 'C0392B')

note(doc, '※ STEP番号を指しながら順番に。STEP2は特に強調')

script(doc, 'カスハラに遭遇したときの対応フローです。5ステップで覚えてください。')
for step, title, detail in [
    ('STEP 1', '不当な言動に気づく',   '怒鳴られているとパニックになりますが、「これはクレームか、カスハラか」と一歩引いて状況を見てください。'),
    ('STEP 2', '一人で抱え込まない',   'すぐに同僚・先輩・リーダーへ声をかける——これが最重要です。一人で対応しようとすると被害が大きくなります。'),
    ('STEP 3', '複数人で対応する',     '必ず2名以上。これが大原則です。'),
    ('STEP 4', '事実を記録する',       '日時・言動の内容を具体的にメモしておく。後で組織として対応するときの証拠になります。'),
    ('STEP 5', '上長・窓口へ報告する', '組織として対応する。ここで終わりではなく、ここからが組織の出番です。'),
]:
    emphasis(doc, f'★ {step}　{title}', indent=0.3)
    script(doc, detail, indent=0.8)

script(doc, '2024年の法改正で、事業者はカスハラへの対策・相談体制・スタッフへのフォローが義務化されました。')
emphasis(doc, '★ 一人で我慢するのは間違っています。カスハラは組織で対応するものです。')


# ============================================================
# スライド 10: 法令遵守 — 主要法令
# ============================================================
slide_header(doc, 10, '法令遵守　―　主要法令', '0D7377')

note(doc, '※ 全法令を細かく読まない。要点だけ触れてスピードを保つ')

script(doc, '最後に法令遵守です。医療現場に関わる主な法令を確認します。')
script(doc, '個人情報保護法——先ほどお話しした通りです。漏洩時は72時間以内の報告義務があります。', indent=0.5)
script(doc, '医療法——医療記録の正確な保管義務と、患者への説明責任が求められます。記録の改ざんは違反です。', indent=0.5)
script(doc, '労働施策総合推進法、いわゆるパワハラ防止法——カスハラへの対策も事業者の義務です。', indent=0.5)
script(doc, '医師法・保健師助産師看護師法——業務上知った情報の守秘義務。退職後も一生続きます。', indent=0.5)
emphasis(doc, '★ 違反した場合——刑事罰、行政処分（免許停止・取り消し）、民事賠償の3つが待っています。')
script(doc, '免許の停止・取り消しは、医療専門職としてのキャリアが終わることを意味します。')


# ============================================================
# スライド 11: 法令遵守 — 相談・報告
# ============================================================
slide_header(doc, 11, '法令遵守　―　迷ったら報告', '0D7377')

emphasis(doc, '★ 一番伝えたいことは——「迷ったら報告してください」ということです。')
script(doc, 'ステップ1——「これはまずいかも？」と思ったら、まず直属の上長へ。', indent=0.5)
script(doc, 'ステップ2——上長に言いにくい場合は、院内の相談窓口を使ってください。', indent=0.5)
script(doc, 'ステップ3——隠さない・後回しにしない。早期報告が被害を最小化します。', indent=0.5)
script(doc, '問題が起きたとき、隠したくなる気持ちはわかります。でも隠すほど被害は広がります。')
emphasis(doc, '★ 早期報告が自分を・患者を・病院を守ります。')
script(doc, 'スライド下部に院内の相談窓口を記載しています。後で記入して周知してください。')


# ============================================================
# スライド 12: まとめ
# ============================================================
slide_header(doc, 12, 'まとめ　―　3つのメッセージ', '1A1A2E')

script(doc, '最後にまとめです。今日お伝えしたことを3つに絞ります。')

for num, color_name, title, desc in [
    ('01', 'ブルー',   '患者情報はあなたが「預かっている」もの',
                      '慎重に、目的の範囲内だけで扱ってください。SNS投稿・私用スマホ撮影・廊下での会話——全部ダメです。'),
    ('02', 'レッド',   'ハラスメントは職場の問題',
                      'カスハラは特に一人で抱え込まないでください。必ず複数人で、組織として対応してください。'),
    ('03', 'ティール', '迷ったら報告',
                      '隠さないことが、自分と患者と病院を守ります。早期報告が最善の対処です。'),
]:
    emphasis(doc, f'★ メッセージ{num}　{title}', indent=0.3)
    script(doc, desc, indent=0.8)

script(doc, 'この3つを持ち帰って、明日からの業務に活かしてください。')


# ============================================================
# スライド 13: セルフチェック
# ============================================================
slide_header(doc, 13, 'セルフチェックリスト', 'F39C12')

note(doc, '※ 最後のスライド。チェックリストを配布資料で確認するよう促す')

script(doc, '最後に、配布資料のセルフチェックリストを使って自己確認をしてください。')
script(doc, '個人情報保護・ハラスメント予防・法令遵守——それぞれチェックボックスがあります。')
emphasis(doc, '★ チェックできない項目が、今日から変えるポイントです。')
script(doc, 'チェックリストは職場に持ち帰って、定期的に振り返る習慣をつけてください。')
script(doc, 'お疲れさまでした。ご質問はありますか？')


# ============================================================
# 所要時間メモ
# ============================================================
spacer(doc, 6)
p_br = doc.add_paragraph()
set_border_bottom(p_br, sz=4)
run_br = p_br.add_run('')
run_br.font.size = Pt(4)
set_run_font(run_br)

p_time = doc.add_paragraph()
set_shading(p_time, 'F5F5F5')
p_time.paragraph_format.left_indent = Cm(0.3)
run_t = p_time.add_run(
    '【所要時間の目安】\n'
    'スライド1〜2：イントロ＋アジェンダ　約1分\n'
    'スライド3〜5：個人情報保護　約6分\n'
    'スライド6〜9：ハラスメント予防（カスハラ重点）　約8分\n'
    'スライド10〜11：法令遵守　約4分\n'
    'スライド12〜13：まとめ＋チェックリスト　約1分\n'
    '合計：約20分　＋　質疑応答10分　＝　計30分'
)
run_t.font.size = Pt(10)
run_t.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
set_run_font(run_t)


# ===== 保存 =====
out = 'C:/Users/yuuna/agens/勉強会/発表原稿_スライド用.docx'
doc.save(out)
print(f'保存完了：{out}')
