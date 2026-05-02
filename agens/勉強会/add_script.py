# -*- coding: utf-8 -*-
"""
配布資料_new.docx に発表原稿セクションを追加する。
各スライドに対応したスピーカーノートを文書末尾に挿入。
"""
import docx, io, sys
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc = docx.Document('配布資料_new.docx')

# ─────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────
def new_p():
    return OxmlElement('w:p')

def set_spacing(p_elem, before_pt=None, after_pt=None):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if before_pt is not None:
        sp.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt is not None:
        sp.set(qn('w:after'), str(int(after_pt * 20)))

def set_indent(p_elem, left_pt=None):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if left_pt is not None:
        ind.set(qn('w:left'), str(int(left_pt * 20)))

def set_shading(p_elem, fill_hex):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def add_run(p_elem, text, bold=False, size_pt=None, color_hex=None, italic=False):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    if size_pt is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        rPr.append(szCs)
    if color_hex:
        col = OxmlElement('w:color')
        col.set(qn('w:val'), color_hex)
        rPr.append(col)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p_elem.append(r)

def append_para(p_elem):
    """Append a paragraph element to the document body."""
    doc.element.body.append(p_elem)

# ─────────────────────────────────────────────
# Paragraph builders for 発表原稿
# ─────────────────────────────────────────────

def make_page_break():
    """改ページ"""
    p = new_p()
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p

def make_main_header(text):
    """発表原稿 大見出し (00スタイル: bold 15pt color=111111, before=10pt after=5pt)"""
    p = new_p()
    set_spacing(p, before_pt=10, after_pt=5)
    add_run(p, text, bold=True, size_pt=15, color_hex='111111')
    return p

def make_slide_label(text):
    """スライドラベル: 黒背景 bold 10pt white (before=4pt after=0 left=11.375pt)"""
    p = new_p()
    set_spacing(p, before_pt=4, after_pt=0)
    set_indent(p, left_pt=11.375)
    set_shading(p, '333333')
    add_run(p, text, bold=True, size_pt=10, color_hex='FFFFFF')
    return p

def make_script_line(text, last=False):
    """発表原稿テキスト: 薄いグレー背景 10pt (before=0 after=0/4pt left=11.375pt)"""
    p = new_p()
    set_spacing(p, before_pt=0, after_pt=4 if last else 0)
    set_indent(p, left_pt=11.375)
    set_shading(p, 'F0F0F0')
    add_run(p, text, size_pt=10)
    return p

def make_cue_line(text):
    """（動作指示）行: グレー背景 + italic 9pt color=666666"""
    p = new_p()
    set_spacing(p, before_pt=0, after_pt=0)
    set_indent(p, left_pt=11.375)
    set_shading(p, 'F0F0F0')
    add_run(p, text, size_pt=9, color_hex='666666', italic=True)
    return p

def add_section(label, lines):
    """スライドラベル＋スクリプト行をまとめて追加。
    lines: list of (text, is_cue) tuples. is_cue=True → 動作指示行（斜体グレー）
    """
    append_para(make_slide_label(label))
    for i, item in enumerate(lines):
        if isinstance(item, tuple):
            text, is_cue = item
        else:
            text, is_cue = item, False
        last = (i == len(lines) - 1)
        if is_cue:
            append_para(make_cue_line(text))
        else:
            append_para(make_script_line(text, last=last))

# ─────────────────────────────────────────────
# 発表原稿 本文
# ─────────────────────────────────────────────

append_para(make_page_break())
append_para(make_main_header('発表原稿'))

# --- 00 アジェンダ ---
add_section('【00　本日のアジェンダ】', [
    ('（配布資料を配りながら）', True),
    'お忙しい中、本日の院内研修にご参加いただきありがとうございます。',
    '今日の研修では「個人情報保護」「ハラスメント予防」「法令遵守」の3つのテーマを取り上げます。',
    'いずれも皆さんの日常業務に直結する内容です。',
    '配布資料をお手元に置いて、ご自身の業務を振り返りながらご参加ください。',
    ('所要時間はおよそ○○分を予定しています。', True),
])

# --- 01 個人情報保護 ---
add_section('【01　個人情報保護 ─ 「個人情報」とは何か？】', [
    'まず「個人情報保護」から確認していきます。',
    '「個人情報」とは、氏名・生年月日・住所など特定の個人を識別できる情報です。',
    '病院では患者さんの診断名・処方内容・検査結果・家族関係なども含まれます。',
    'これらは皆さんが業務上「預かっている」情報です。患者さんの同意なく外部へ漏らすことは法律で禁止されています。',
])

add_section('【01　患者情報取り扱い 5原則】', [
    '患者情報を安全に扱うための5つの原則を確認します。',
    ('（スライドを指しながら各原則を読み上げ、1〜2文のコメントを添える）', True),
    '日常業務の中で、これらの原則を意識する習慣をつけてください。',
])

add_section('【01　やってはいけないこと NG ─ SNS投稿について特に注意】', [
    'やってはいけないことを一覧で確認します。',
    'たとえば、私用スマートフォンで患者情報を撮影すること。これは絶対にNGです。個人のスマホはセキュリティが保証されていません。',
    'SNSへの患者関連の投稿も、患者が特定できない内容であっても一切禁止です。「バレないから大丈夫」は通用しません。',
    '新たに追加したルールがあります。制服姿や院内の写真・動画をプライベートSNSに投稿することも禁止です。',
    '制服や背景から勤務先が特定されるだけでなく、写り込んだ書類や画面が患者情報の漏洩になる場合があります。',
    'さらに、投稿した位置情報や行動パターンからスタッフ自身がストーカー被害に遭うケースも実際に起きています。',
    '「自分の写真だから問題ない」という認識は誤りです。今日からSNS投稿の際は必ず確認してください。',
])

add_section('【01　違反した場合のリスク】', [
    '万一違反した場合、個人として「1年以下の懲役または50万円以下の罰金」という刑事罰を受ける可能性があります。',
    '加えて、病院全体の社会的信用の失墜、患者さんとの信頼関係の崩壊につながります。',
    '「知らなかった」では済まされません。今日の研修でしっかり理解していただくことが最大の予防策です。',
])

# --- 02 カスハラ ---
add_section('【02　ハラスメントの主な種類】', [
    '続いてハラスメント予防です。',
    '職場のハラスメントにはパワハラ・セクハラ・マタハラなど様々な種類がありますが、',
    '今日は特に病院現場で増加している「カスタマーハラスメント（カスハラ）」と「パワーハラスメント（パワハラ）」を重点的に解説します。',
])

add_section('【02　カスタマーハラスメント（カスハラ）─ 定義】', [
    'カスハラとは、患者さんやそのご家族が行う「社会通念上不相当な言動」で、スタッフの就業環境を害するものです。',
    '「クレーム」と「カスハラ」は違います。正当な不満・要望は丁寧に対応すべきですが、',
    '怒鳴る・脅す・長時間にわたって拘束するといった行為はカスハラに当たります。',
    'スタッフが「患者だから仕方ない」と一人で我慢する必要はありません。',
])

add_section('【02　病院で起こりやすいカスハラの具体例】', [
    ('（スライドの具体例を指しながら）', True),
    'こうした言動がカスハラに該当します。',
    '「患者だから何を言ってもいい」ということはありません。スタッフには安全に働く権利があります。',
])

add_section('【02　カスハラへの対応フロー（スタッフ向け）】', [
    'カスハラを受けたときの対応は5ステップです。',
    'まず「これはクレームか、カスハラか」を判断する。次に、絶対に一人で抱え込まないこと。',
    'すぐに同僚・先輩に声をかけ、必ず2名以上で対応してください。単独対応は禁止です。',
    '対応後は、日時・言動の内容を具体的にメモし、上長・相談窓口へ必ず報告してください。',
    '記録を残すことは、組織として対応するためにも、皆さん自身を守るためにも不可欠です。',
])

add_section('【02　組織の責任】', [
    '2024年の労働施策総合推進法の指針改正により、事業者にはカスハラ対策・相談体制の整備が義務化されました。',
    'スタッフの皆さんが一人で我慢する必要はありません。困ったときは必ず組織として対応します。',
    '何かあれば、遠慮なく上長または相談窓口に声をかけてください。',
])

# --- 02-2 パワハラ ---
add_section('【02-2　パワーハラスメント（パワハラ）─ 定義・6類型】', [
    '次にパワーハラスメントです。パワハラは厚生労働省が定める3つの要件をすべて満たす行為です。',
    '6つの類型の中で特に注意してほしいのが「精神的な攻撃」と「人間関係からの切り離し」です。',
    '怒鳴る・侮辱的な言葉を使う・無視する・仲間外しにする行為は典型的なパワハラです。',
    '「指導のつもりだった」「本人のためを思って言った」という認識があっても、',
    '相手が精神的なダメージを受け、就業環境が害されればパワハラと判断されます。',
])

add_section('【02-2　パワハラをしないために気をつけること】', [
    '特に意識してほしいのは「人格ではなく行動を指摘する」という点です。',
    '「あなたはダメだ」という言い方は厳禁です。「この書類の○○の部分を直してほしい」など、具体的な行動を指摘してください。',
    '注意・指導は必ず個別・非公開の場で行ってください。人前での叱責は相手の尊厳を傷つけ、パワハラと受け取られます。',
    '「これくらい当たり前」という感覚も危険です。自分の「当たり前」が、相手には過大な負担である場合があります。',
    '自分の言動を振り返る習慣を持つこと、そして迷ったら一人で判断せず相談窓口に声をかけてください。',
])

# --- 03 法令遵守 ---
add_section('【03　法令遵守（コンプライアンス）】', [
    '最後に法令遵守です。医療現場には守るべき法律が数多くあります。',
    ('（スライドの主要法令を指しながら）', True),
    'これらは「知らなかった」では守れません。代表的なものを今日改めて確認してください。',
    '特に守秘義務は退職後も続きます。前の職場の情報を転職先で話すことも違反になります。注意してください。',
])

add_section('【03　迷ったとき・気づいたときは「報告・相談」が最優先】', [
    '迷ったときのルールはシンプルです。「これはまずいかも？」と思ったら、まず直属の上長に相談してください。',
    '隠したり後回しにすることが最大のリスクです。',
    '早期に報告することで、多くの問題は小さなうちに解決できます。',
    '院内相談窓口もありますので、上長への相談が難しい場合は迷わず利用してください。',
])

# --- セルフチェックリスト ---
add_section('【✓　セルフチェックリスト】', [
    '配布資料のチェックリストを使って、ご自身の行動を振り返ってください。',
    ('（参加者に記入時間を1〜2分設ける）', True),
    'チェックできない項目があれば、それが今日からの改善ポイントです。ぜひ正直に振り返ってみてください。',
    'チェックリストは持ち帰って、日常業務の中で定期的に確認するようにしてください。',
])

# --- まとめ ---
add_section('【まとめ：3つのメッセージ】', [
    '本日の研修の3つのメッセージをもう一度確認します。',
    '1つ目。患者情報はあなたが「預かっている」ものです。慎重に扱ってください。',
    '2つ目。ハラスメントは職場全体の問題です。一人で抱え込まず、組織で対応しましょう。',
    '3つ目。迷ったら必ず報告する。隠さないことが、自分と患者さんと病院を守ります。',
    'この3つを、今日から日常業務の中で意識してください。',
    '本日の研修へのご参加、ありがとうございました。質問のある方はこの後でお声がけください。',
    ('（質疑応答）', True),
])

doc.save('配布資料_final.docx')
print('Done. 発表原稿 added.')
