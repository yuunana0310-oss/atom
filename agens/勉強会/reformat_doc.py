# -*- coding: utf-8 -*-
"""
配布資料_new.docx のパワハラセクション（逆順になっている）を削除して正しい順序で再挿入する。
SNS・チェックリストはすでに正しいので触らない。
"""
import docx, io, sys
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc = docx.Document('配布資料_new.docx')

# ─────────────────────────────────────────────
# STEP 1: Delete wrong-order パワハラ paragraphs
# ─────────────────────────────────────────────
DELETE_MARKERS = [
    '02-2　パワーハラスメント（パワハラ）重点解説',
    '職場において行われる①優越的な',
    '②業務上必要かつ相当な範囲を超えた',
    '③労働者の就業環境が害される',
    'パワハラの6類型',
    '① 身体的な攻撃',
    '② 精神的な攻撃',
    '③ 人間関係からの切り離し',
    '④ 過大な要求',
    '⑤ 過小な要求',
    '⑥ 個の侵害',
    'パワハラをしないために気をつけること',
    '指導と叱責を区別する',
    '人格ではなく行動を指摘する',
    '人前での叱責を避ける',
    '業務指示は適切な範囲内で',
    'これくらい当然',
    '無視・孤立化をしない',
    '自分の言動を振り返る習慣',
    '迷ったら相談窓口へ',
    '【注意】  指導・教育の範囲内',
    # Old format ones still in original
    '◎ 指導と叱責を区別する',
    '◎ 人格ではなく行動を指摘する',
    '◎ 人前での叱責を避ける',
    '◎ 業務指示は適切な範囲内で',
    '◎ 無視・孤立化をしない',
    '◎ 自分の言動を振り返る習慣',
    '◎ 迷ったら相談窓口へ',
    '【注意】指導・教育の範囲内',
]

paras_to_del = []
for para in doc.paragraphs:
    for marker in DELETE_MARKERS:
        if marker in para.text:
            paras_to_del.append(para._p)
            break

for p in paras_to_del:
    try:
        p.getparent().remove(p)
    except Exception:
        pass

print(f'Deleted {len(paras_to_del)} paragraphs')

# ─────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────
def new_p():
    return OxmlElement('w:p')

def set_spacing(p_elem, before_pt=None, after_pt=None):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if before_pt is not None:
        sp.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt is not None:
        sp.set(qn('w:after'), str(int(after_pt * 20)))

def set_indent(p_elem, left_pt=None):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if left_pt is not None:
        ind.set(qn('w:left'), str(int(left_pt * 20)))

def set_shading_black(p_elem):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '111111')
    pPr.append(shd)

def add_run(p_elem, text, bold=False, size_pt=None, color_hex=None):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if size_pt is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        rPr.append(szCs)
    if color_hex:
        col = OxmlElement('w:color')
        col.set(qn('w:val'), color_hex)
        rPr.append(col)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p_elem.append(r)

# ─────────────────────────────────────────────
# Paragraph factory functions (matching existing styles)
# ─────────────────────────────────────────────

def make_section_header_14(text):
    """02スタイル: bold 14pt color=111111, space_before=10pt after=5pt (original: 177800=14pt EMU, 127000/63500 twip)"""
    p = new_p()
    set_spacing(p, before_pt=10, after_pt=5)
    add_run(p, text, bold=True, size_pt=14, color_hex='111111')
    return p

def make_definition_first(label, text):
    """【定義】一行目: space_before=4pt after=0, left=11.375pt, label=9.5pt bold, text=10pt"""
    p = new_p()
    set_spacing(p, before_pt=4, after_pt=0)
    set_indent(p, left_pt=11.375)
    add_run(p, label, bold=True, size_pt=9.5)
    add_run(p, text, size_pt=10)
    return p

def make_definition_line(text, last=False):
    """定義続き行"""
    p = new_p()
    set_spacing(p, before_pt=0, after_pt=6 if last else 0)
    set_indent(p, left_pt=11.375)
    add_run(p, text, size_pt=10)
    return p

def make_subsection_header(text):
    """中見出し: bold, space_before=8pt after=4pt — matches 病院で起こりやすい/カスハラへの対応フロー style"""
    p = new_p()
    set_spacing(p, before_pt=8, after_pt=4)
    add_run(p, text, bold=True)
    return p

def make_ng_item(label, text):
    """NG/類型リスト: sp=1pt, left=22.75pt (288290 EMU), label bold 10pt + text 10pt"""
    p = new_p()
    set_spacing(p, before_pt=1, after_pt=1)
    set_indent(p, left_pt=22.75)
    add_run(p, label, bold=True, size_pt=10)
    add_run(p, text, size_pt=10)
    return p

def make_black_header(text):
    """黒背景タイトル行: shading, white bold 10pt, space_before=4pt after=0, left=11.375pt"""
    p = new_p()
    set_spacing(p, before_pt=4, after_pt=0)
    set_indent(p, left_pt=11.375)
    set_shading_black(p)
    add_run(p, text, bold=True, size_pt=10, color_hex='FFFFFF')
    return p

def make_black_line(text, last=False):
    """黒背景コンテンツ行"""
    p = new_p()
    set_spacing(p, before_pt=0, after_pt=6 if last else 0)
    set_indent(p, left_pt=11.375)
    set_shading_black(p)
    add_run(p, text, size_pt=10, color_hex='FFFFFF')
    return p

def make_note(label, text, last=False):
    """【覚えておこう】スタイル: label=9.5pt bold, text=10pt, space_before=4pt after=0/6pt, left=11.375pt"""
    p = new_p()
    set_spacing(p, before_pt=4, after_pt=6 if last else 0)
    set_indent(p, left_pt=11.375)
    add_run(p, label, bold=True, size_pt=9.5)
    add_run(p, text, size_pt=10)
    return p

# ─────────────────────────────────────────────
# STEP 2: Build パワハラ section in correct order
# ─────────────────────────────────────────────
ref_hoshu = None
for para in doc.paragraphs:
    if '法令遵守（コンプライアンス）' in para.text and '03' in para.text:
        ref_hoshu = para._p
        break
assert ref_hoshu is not None, '03法令遵守 not found'

pw_paras = []

# ① Section header
pw_paras.append(make_section_header_14('02-2　パワーハラスメント（パワハラ）重点解説'))

# ② Definition block (3 lines)
pw_paras.append(make_definition_first('【定義】  ', '職場において行われる①優越的な関係を背景とした言動であって、'))
pw_paras.append(make_definition_line('②業務上必要かつ相当な範囲を超えたものにより、'))
pw_paras.append(make_definition_line('③労働者の就業環境が害されるもの。（厚生労働省「パワーハラスメント防止指針」より）', last=True))

# ③ 6類型 sub-header + items
pw_paras.append(make_subsection_header('パワハラの6類型（厚労省定義）'))
types = [
    ('① 身体的な攻撃　　　', '— 暴行・傷害'),
    ('② 精神的な攻撃　　　', '— 脅迫・名誉毀損・侮辱・ひどい暴言'),
    ('③ 人間関係からの切り離し　', '— 隔離・仲間外し・無視'),
    ('④ 過大な要求　　　　', '— 業務上明らかに不要・不可能なことの強制'),
    ('⑤ 過小な要求　　　　', '— 能力・経験を無視した程度の低い業務のみ命じること'),
    ('⑥ 個の侵害　　　　　', '— 私的なことへの過度な立ち入り'),
]
for label, desc in types:
    pw_paras.append(make_ng_item(label, desc))

# ④ パワハラ防止 black box
pw_paras.append(make_black_header('パワハラをしないために気をつけること'))
prevention = [
    '指導と叱責を区別する — 感情的にならず、事実に基づいて具体的に伝える',
    '人格ではなく行動を指摘する — 「あなたはダメだ」でなく「この点を改善してほしい」',
    '人前での叱責を避ける — 注意・指導は個別・非公開の場で行う',
    '業務指示は適切な範囲内で — 業務外の私的な用事や無理な残業を強要しない',
    '「これくらい当然」思考に注意 — 世代・経験の違いで「当然」は人によって異なる',
    '無視・孤立化をしない — 挨拶・情報共有から意図的に外すことはパワハラになりうる',
    '自分の言動を振り返る習慣を持つ — 「相手がどう受け取ったか」を常に意識する',
    '迷ったら相談窓口へ — パワハラか判断できないときは一人で抱え込まず報告する',
]
for i, pv in enumerate(prevention):
    pw_paras.append(make_black_line(pv, last=(i == len(prevention) - 1)))

# ⑤ 注意 note
pw_paras.append(make_note('【注意】  ',
    '指導・教育の範囲内であっても、継続的・執拗に繰り返される場合はパワハラと判断されることがあります。',
    last=True))

# ─── Insert strategy ───
# addprevious(p) puts p immediately before ref_hoshu.
# Each call pushes p to the slot right before ref_hoshu, shifting previous ones further away.
# To get order [p0, p1, ..., pN, ref_hoshu], insert in FORWARD order (p0 first, pN last).
# After all inserts: p0 is furthest from ref, pN is nearest — reading gives p0...pN, ref. ✓
for p in pw_paras:
    ref_hoshu.addprevious(p)

print(f'Inserted {len(pw_paras)} パワハラ paragraphs in correct order')

# ─────────────────────────────────────────────
# STEP 3: Fix checklist items formatting to match existing style
# (existing: sp_b=12700 sp_a=12700 ind=360045 — 1pt spacing, 28.35pt indent)
# ─────────────────────────────────────────────
for para in doc.paragraphs:
    if para.text in (
        '□  制服・院内が特定できる写真・動画をSNSに投稿していない',
        '□  部下・後輩への指導は感情的にならず、事実に基づいて行っている',
        '□  人格否定・無視・過大または過小な業務指示をしていない',
    ):
        pf = para.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.left_indent = Pt(28.35)

print('Checklist formatting fixed')

doc.save('配布資料_new.docx')
print('Saved: 配布資料_new.docx')
