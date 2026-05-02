from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ余白設定
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_para(text, bold=False, size=11, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

# ===== ヘッダー =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('自主トレーニングガイド')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
r = p.add_run('骨盤底筋（肛門括約筋）')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
r = p.add_run('セルフエクササイズ')
r.font.size = Pt(22)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(12)
r = p.add_run('理学療法士が指導する、毎日できるトレーニングプログラムです')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('─' * 60).paragraph_format.space_after = Pt(10)

# ===== 目的ボックス =====
p = add_para('🎯 このトレーニングで改善できること', bold=True, size=12, space_before=4, space_after=6)
items = ['✔ 便もれ・尿もれの予防', '✔ 術後の括約筋機能回復', '✔ 排便コントロール力の向上', '✔ 産後の骨盤底筋回復']
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()

# ===== 姿勢のステップアップ =====
p = add_para('▌ 姿勢のステップアップ', bold=True, size=13, space_before=8, space_after=4)
p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

p = add_para('慣れてきたら少しずつ姿勢を変えていきましょう', size=11, space_before=0, space_after=6)

tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
cells = tbl.rows[0].cells
cells[0].text = '仰向け\n（膝立て位）\n★まずここから'
cells[1].text = '→'
cells[2].text = '座位\n（椅子）\n慣れてきたら'
cells[3].text = '→'
cells[4].text = '立位\n（立った状態）\nさらに慣れたら'

set_cell_bg(cells[0], '1a1a1a')
for run in cells[0].paragraphs[0].runs:
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for i, cell in enumerate(cells):
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)

p0 = cells[0].paragraphs[0]
p0.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
p0.runs[0].bold = True
set_cell_bg(cells[0], '1a1a1a')

doc.add_paragraph()

# ===== 注意事項 =====
p = add_para('▌ はじめる前に確認しましょう', bold=True, size=13, space_before=8, space_after=4)

cautions = [
    ('🛏️', '最初は仰向けで行う', '膝を立てて寝た状態がもっとも力を入れやすい姿勢です'),
    ('🌬️', '呼吸を止めない', 'ゆっくり呼吸しながら行いましょう'),
    ('🎯', '意識するのは肛門だけ', 'お腹・太もも・お尻は脱力したまま'),
    ('🛑', '痛みがあれば中止', '担当の理学療法士へご相談ください'),
]

tbl2 = doc.add_table(rows=2, cols=2)
tbl2.style = 'Table Grid'
for idx, (icon, title, desc) in enumerate(cautions):
    row = idx // 2
    col = idx % 2
    cell = tbl2.rows[row].cells[col]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{icon} {title}\n')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    set_cell_bg(cell, 'f7f7f7')

doc.add_paragraph()

# ===== エクササイズ =====
p = add_para('▌ エクササイズの内容', bold=True, size=13, space_before=8, space_after=4)

exercises = [
    {
        'num': '1',
        'title': 'ゆっくり締める（持続収縮）',
        'subtitle': '遅筋を鍛える／安静時の括約筋持続力を高める',
        'steps': [
            '仰向けに寝て、両膝を立てる（膝立て位）',
            '全身の力を抜き、お腹・太もも・お尻をリラックスさせる',
            '肛門を「上に引き上げるように」ゆっくり締める',
            '⏱ 5〜10秒 キープ',
            'ゆっくり力を抜いて、同じ時間休む',
        ],
        'tip': 'エレベーターが1階→2階→3階と上がるイメージで段階的に締める',
        'reps': '× 10回 ／ 1セット',
    },
    {
        'num': '2',
        'title': '素早く締める（速収縮）',
        'subtitle': '速筋を鍛える／急な腹圧（咳・くしゃみ）への対応力を高める',
        'steps': [
            '仰向けに寝て、両膝を立てる（①と同じ姿勢）',
            '肛門をパッと素早く締める',
            '⏱ 1秒 締める → 1秒 抜く',
            'リズムよく繰り返す',
        ],
        'tip': '「スイッチをON・OFFするイメージ」でテンポよく行う',
        'reps': '× 10回 ／ 1セット',
    },
    {
        'num': '3',
        'title': '段階的収縮（グレーデッドエクササイズ）',
        'subtitle': '随意的なコントロール能力を高める／①②に慣れてから行う',
        'steps': [
            '仰向けに寝て、両膝を立てる',
            '20%の力でゆっくり締める',
            '50%まで徐々に力を強くする',
            '100%まで締めて ⏱ 3秒 キープ',
            '段階的にゆっくり力を抜いていく',
            '完全にリラックスして次へ',
        ],
        'tip': '力の「強さ」を意識することで随意的な筋コントロールが上達する',
        'reps': '× 5回 ／ 1セット',
    },
]

for ex in exercises:
    # ヘッダー行
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, '1a1a1a')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'【{ex["num"]}】 {ex["title"]}')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = cell.add_paragraph(ex['subtitle'])
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)

    # ステップ
    for i, step in enumerate(ex['steps'], 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.runs[0].font.size = Pt(11)

    # tip
    p = doc.add_paragraph(f'💡 {ex["tip"]}')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # reps
    p = doc.add_paragraph(ex['reps'])
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(11)

# ===== スケジュール =====
p = add_para('▌ 1日のトレーニングスケジュール', bold=True, size=13, space_before=8, space_after=6)

tbl3 = doc.add_table(rows=2, cols=3)
tbl3.style = 'Table Grid'
schedule = [
    ('🌅 朝（起床後）', '①持続収縮 10回\n②速収縮 10回'),
    ('☀️ 昼（食後など）', '①持続収縮 10回'),
    ('🌙 夜（就寝前）', '①持続収縮 10回\n③段階的収縮 5回'),
]
for col, (time, items) in enumerate(schedule):
    cell = tbl3.rows[0].cells[col]
    set_cell_bg(cell, 'f0f0f0')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(time)
    r.bold = True
    r.font.size = Pt(11)

    cell2 = tbl3.rows[1].cells[col]
    cell2.text = ''
    p2 = cell2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(items)
    r2.font.size = Pt(10)

# 合計
p = add_para('合計目安：1日 3セット ／ 約10〜15分　｜　姿勢はすべて仰向けで行いましょう',
             bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=10)

# ===== NG表 =====
p = add_para('▌ よくある間違い（NG例と正しい方法）', bold=True, size=13, space_before=8, space_after=6)

ng_data = [
    ('✕ お腹（腹筋）に力が入る', '○ 腹筋はリラックスしたまま'),
    ('✕ 息を止めて行う', '○ ゆっくり呼吸しながら続ける'),
    ('✕ 太ももを内側に絞る', '○ 太ももは脱力したまま'),
    ('✕ 急いでセット数をこなす', '○ 1回1回丁寧に締める・緩める'),
    ('✕ 痛みを我慢して続ける', '○ 痛みがあれば即中止・相談する'),
]

tbl4 = doc.add_table(rows=len(ng_data)+1, cols=2)
tbl4.style = 'Table Grid'
# ヘッダー
for col, text in enumerate(['✕ やりがちなNG', '○ 正しい方法']):
    cell = tbl4.rows[0].cells[col]
    set_cell_bg(cell, '1a1a1a')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for i, (ng, ok) in enumerate(ng_data, 1):
    tbl4.rows[i].cells[0].text = ng
    tbl4.rows[i].cells[1].text = ok
    tbl4.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    tbl4.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    for j in range(2):
        tbl4.rows[i].cells[j].paragraphs[0].paragraph_format.space_before = Pt(3)
        tbl4.rows[i].cells[j].paragraphs[0].paragraph_format.space_after = Pt(3)
    if i % 2 == 0:
        set_cell_bg(tbl4.rows[i].cells[0], 'f7f7f7')
        set_cell_bg(tbl4.rows[i].cells[1], 'f7f7f7')

doc.add_paragraph()

# ===== 継続のコツ =====
p = add_para('▌ 継続するためのヒント', bold=True, size=13, space_before=8, space_after=6)

tips = [
    ('📺 「ながら運動」でOK', '寝たままテレビを見ながら、就寝前の布団の中など日常に組み込みましょう'),
    ('📅 記録をつける', 'カレンダーや手帳に✓を入れると継続しやすくなります'),
    ('⏳ 効果まで4〜8週間', '焦らず続けることが大切です。変化が出るまで継続しましょう'),
    ('⬆️ 慣れたら姿勢を変える', '仰向けで楽にできたら、座位・立位へとステップアップしましょう'),
]

tbl5 = doc.add_table(rows=2, cols=2)
tbl5.style = 'Table Grid'
for idx, (title, desc) in enumerate(tips):
    row = idx // 2
    col = idx % 2
    cell = tbl5.rows[row].cells[col]
    set_cell_bg(cell, 'f0f0f0')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + '\n')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

# ===== 相談のタイミング =====
p = add_para('▶ このような場合はすぐに担当者へご相談ください', bold=True, size=12, space_before=8, space_after=6)

warnings = [
    'トレーニング中・後に痛みや違和感が出た',
    '症状が悪化したと感じる',
    'やり方がわからなくなった',
    '1週間以上継続できていない',
]
for w in warnings:
    p = doc.add_paragraph(f'▶ {w}', style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph('─' * 60).paragraph_format.space_before = Pt(8)

# ===== フッター =====
tbl6 = doc.add_table(rows=2, cols=3)
tbl6.style = 'Table Grid'
labels = ['施設名', '担当理学療法士', '連絡先']
for col, label in enumerate(labels):
    cell = tbl6.rows[0].cells[col]
    set_cell_bg(cell, '1a1a1a')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    cell2 = tbl6.rows[1].cells[col]
    cell2.text = ''
    cell2.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell2.paragraphs[0].paragraph_format.space_after = Pt(16)

p = add_para('作成日：2026年3月', size=9, align=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=4, space_after=0)
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

output_path = r'C:\Users\yuuna\agens\肛門括約筋トレーニング_配布資料.docx'
doc.save(output_path)
print(f'保存完了: {output_path}')
