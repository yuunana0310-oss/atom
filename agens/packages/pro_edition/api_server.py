import os
import sys
import json
import urllib.request
import urllib.error
from http.server import SimpleHTTPRequestHandler, HTTPServer
import glob
from pathlib import Path
import io
try:
    import docx
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    docx = None

# RAGエンジンを読み込むためパスを追加
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PRO_DB_DIR = os.path.join(BASE_DIR, 'knowledge', 'pro_db')
AGENTS_DIR = os.path.join(BASE_DIR, 'agents')
UI_DIR = os.path.join(BASE_DIR, 'ui')

sys.path.insert(0, PRO_DB_DIR)

# --- Agent Routing Logic ---
AGENT_PROFILES = {
    "document_drafter": {"name": "書類作成担当", "keywords": ["作成", "書類", "契約", "雛形", "書面", "文面", "ドラフト", "規程", "草案", "お知らせ"]},
    "sns_strategist": {"name": "SNS・集客担当", "keywords": ["SNS", "集客", "投稿", "マーケティング", "X", "Twitter", "ブログ", "告知", "キャンペーン", "インフルエンサー"]},
    "strategy_adviser": {"name": "経営・戦略アドバイザー", "keywords": ["経営", "戦略", "採用", "利益", "売上", "ビジネスモデル", "方針", "課題", "事業"]},
    "client_communicator": {"name": "顧客対応担当", "keywords": ["連絡", "メール", "クレーム", "対応", "案内", "お詫び", "返信", "謝罪", "顧客"]},
    "researcher": {"name": "リサーチ担当", "keywords": ["調査", "ニュース", "最新", "判例", "動向", "教えて", "違い", "改正"]},
    "expert_assistant": {"name": "実務アシスタント", "keywords": ["手続き", "実務", "判断", "申請", "届出", "方法", "どうすれば", "対策", "制度"]}
}

def determine_best_agent(query):
    """質問内容から最適なエージェントをスコアリングで自動判定する"""
    scores = {agent_id: 0 for agent_id in AGENT_PROFILES.keys()}
    
    for agent_id, data in AGENT_PROFILES.items():
        for kw in data["keywords"]:
            if kw in query:
                scores[agent_id] += 1
                
    best_agent_id = max(scores, key=scores.get)
    if scores[best_agent_id] == 0:
        best_agent_id = "expert_assistant" # デフォルト
        
    return best_agent_id, AGENT_PROFILES[best_agent_id]["name"]

# 遅延ロード用のグローバル変数
global_rag = None
rag_initialized = False

def get_rag():
    global global_rag, rag_initialized
    if not rag_initialized:
        print("[SYSTEM] RAGエンジンを初回ロード中...（数十秒かかる場合があります）")
        try:
            from pro_rag_engine import ProRAG
            global_rag = ProRAG()
            print("[SYSTEM] RAGエンジンロード完了")
        except Exception as e:
            print(f"[ERROR] RAGエンジンの初期化に失敗しました: {e}")
            global_rag = None
        rag_initialized = True
    return global_rag

def get_agent_prompt(agent_name):
    """エージェントのMDファイルを探して読み込む"""
    pattern = os.path.join(AGENTS_DIR, "*.md")
    for filepath in glob.glob(pattern):
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            if agent_name.lower() in filepath.lower() or agent_name in content.split('\n')[0]:
                return content
    return "あなたは優秀なアシスタントです。"

def generate_master_prompt(agent_name, query):
    """RAGの結果とエージェント設定を合体した最強のプロンプトを作る"""
    base_prompt = get_agent_prompt(agent_name)
    
    # RAGで1.6万件から関連データを検索
    context = ""
    raw_results = []
    rag = get_rag()
    
    if rag:
        raw_results = rag.search(query, n_results=3)
        context = rag.format_context(raw_results)
    else:
        context = "[ERROR] 知能データベースに接続できませんでした。"

    # ニュース要約の読み込み
    news_context = ""
    news_file = os.path.join(UI_DIR, "news_briefing.json")
    if os.path.exists(news_file):
        try:
            with open(news_file, 'r', encoding='utf-8') as f:
                news_data = json.load(f)
                news_context = "【本日の業界ニュース要約】\n"
                for item in news_data:
                    news_context += f"- {item['title']}\n  概要: {item['summary']}\n"
        except Exception:
            pass

    master_prompt = f"""{base_prompt}

--------------------------------------------------
【提供されたコンテキスト情報】
以下の情報は、あなたのためにシステムが自動収集した「RAGデータベースからの抽出知識」および「本日最新の業界ニュース」です。
ユーザーからの指示（タスク）を遂行する上で、関連するものを最大限に活用してください。

{context}

{news_context}

--------------------------------------------------
【あなたの任務と振る舞い方（重要）】
CEO（ユーザー）からの指示に対して「直接的」かつ「臨機応変」に実行してください。

1. もし単なる質問であれば、提供された知識を元に専門家として簡潔に回答してください。
2. もし「〜を作成して」「ブログにして」「翻訳して」といった【作業・タスク】の指示であれば、コンサルタントとしての長文の前置きや挨拶は一切省き、直ちにその作業結果（成果物）のみを出力してください。
3. あなたはただの辞書ではなく有能なアシスタントです。言われたことを文字通りに行うのではなく、プロとして最も適したフォーマットや表現を思考して実行してください。

【CEOからの指示・タスク内容】
{query}
"""
    return master_prompt, raw_results

def call_llm_api(api_key, provider, messages, system_instruction=""):
    """メッセージ履歴（会話メモリ）に対応したAPIコール"""
    if provider == 'openai':
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        # システムプロンプトを先頭に挿入
        full_messages = [{"role": "system", "content": system_instruction}] + messages
        data = {
            "model": "gpt-4o",
            "messages": full_messages,
            "temperature": 0.7
        }
    elif provider == 'gemini':
        url = f"https://generativelanguage.googleapis.com/v1/models/gemini-2.5-flash-lite:generateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
        # Gemini形式に変換（システムプロンプトをメッセージとして統合）
        gemini_contents = []
        # システムプロンプトを最初のユーザー発言の前に「指示」として差し込む
        gemini_contents.append({"role": "user", "parts": [{"text": f"【指示・設定】\n{system_instruction}\n\n上記を踏まえて回答してください。"}]})
        gemini_contents.append({"role": "model", "parts": [{"text": "了解しました。指示された役割と人格、専門知識を完璧に理解しました。これより業務を開始します。"}]})

        for m in messages:
            role = "user" if m['role'] == 'user' else "model"
            gemini_contents.append({"role": role, "parts": [{"text": m['content']}]})
        
        data = {
            "contents": gemini_contents
        }
    else: # anthropic
        url = "https://api.anthropic.com/v1/messages"
        headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        }
        data = {
            "model": "claude-3-5-sonnet-20241022",
            "max_tokens": 4000,
            "system": system_instruction,
            "messages": messages # Anthropicはrole: assistant
        }

    req = urllib.request.Request(url, data=json.dumps(data).encode('utf-8'), headers=headers, method='POST')
    try:
        with urllib.request.urlopen(req) as response:
            res_body = response.read().decode('utf-8')
            res_json = json.loads(res_body)
            if provider == 'openai':
                return res_json['choices'][0]['message']['content']
            elif provider == 'gemini':
                return res_json['candidates'][0]['content']['parts'][0]['text']
            else:
                return res_json['content'][0]['text']
    except urllib.error.URLError as e:
        error_msg = e.read().decode('utf-8') if hasattr(e, 'read') else str(e)
        return f"[API通信エラー]\n原因: {error_msg}"

class UltimateAPIHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=UI_DIR, **kwargs)

    def do_POST(self):
        if self.path == '/api/generate_prompt':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            req = json.loads(post_data.decode('utf-8'))
            
            agent_name = req.get('agent_name', 'auto')
            if agent_name == 'auto':
                agent_id, agent_name = determine_best_agent(req['query'])
            
            master_prompt, raw_results = generate_master_prompt(agent_name, req['query'])
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({
                'prompt': master_prompt, 
                'sources': raw_results,
                'routed_agent': agent_name
            }).encode('utf-8'))
            
        elif self.path == '/api/chat':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            req = json.loads(post_data.decode('utf-8'))
            
            # APIキー確認
            api_key = req.get('api_key', '').strip()
            if not api_key:
                self.send_response(400)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'APIキーが設定されていません。'}).encode('utf-8'))
                return

            agent_name = req.get('agent_name', 'auto')
            query = req.get('query', '')
            history = req.get('history', []) # UIから送られてくる会話履歴

            if agent_name == 'auto':
                agent_id, agent_name = determine_best_agent(query)

            # RAG知識の取得（最新のクエリに基づいて検索）
            _, raw_results = generate_master_prompt(agent_name, query)
            
            # システム指示の構築（知能の核）
            knowledge_text = "\n".join([f"・{r['text']}" for r in raw_results])
            system_instruction = f"""你是 AGENS Ultimate Intelligence Hub 的エリートAI社員、担当は「{agent_name}」です。
以下の専門知識（ナレッジ）をバックボーンに持っています。

【提供された専門知識】
{knowledge_text}

【行動指針】
1. 専門家として、提供された知識を優先的に参照し、誠実に回答してください。
2. 作業依頼（作成・要約・翻訳等）の場合は、前置きを省き、即座に成果物を出力してください。
3. 会話の流れ（履歴）を理解し、一貫性のある対話を行ってください。
"""
            
            # メッセージリストの作成
            # historyがある場合はそれを使用し、最新のqueryを追加
            messages = history + [{"role": "user", "content": query}]
            
            reply = call_llm_api(api_key, req.get('provider', 'anthropic'), messages, system_instruction)
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({
                'reply': reply, 
                'sources': raw_results,
                'routed_agent': agent_name
            }).encode('utf-8'))
            
        elif self.path == '/api/search':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            req = json.loads(post_data.decode('utf-8'))
            query = req.get('query', '').strip()
            
            if not query:
                self.send_response(400)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': '検索クエリが空です。'}).encode('utf-8'))
                return
                
            engine = get_rag()
            raw_docs = engine.search(query, n_results=5) # 5件ほど取得
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'results': raw_docs}).encode('utf-8'))
            
        elif self.path == '/api/refresh_news':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            req = json.loads(post_data.decode('utf-8'))
            self.handle_refresh_news(req)

        elif self.path == '/api/export_docx':
            if docx is None:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'python-docx がインストールされていません。'}).encode('utf-8'))
                return

            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            req = json.loads(post_data.decode('utf-8'))
            text = req.get('text', '').strip()
            
            if not text:
                self.send_response(400)
                self.end_headers()
                return

            # Wordドキュメント作成
            doc = docx.Document()
            
            # 全体の明朝体設定 (日本語対応)
            style = doc.styles['Normal']
            font = style.font
            font.name = 'MS Mincho'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
            font.size = Pt(11)

            # タイトル（1行目を見出しにする）
            lines = text.splitlines()
            if lines:
                h = doc.add_heading(lines[0], 0)
                for run in h.runs:
                    run.font.name = 'MS Gothic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
                
                # 残り
                doc.add_paragraph('\n'.join(lines[1:]))

            # メモリ上へ保存して返す
            f = io.BytesIO()
            doc.save(f)
            f.seek(0)
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', 'attachment; filename="agens_output.docx"')
            self.end_headers()
            self.wfile.write(f.read())
            
        else:
            self.send_response(404)
            self.end_headers()

    def handle_refresh_news(self, data):
        """サーバー側で最新ニュースを取得し、AI要約を実行する"""
        api_key = data.get('api_key')
        
        try:
            import subprocess
            # 1. ニュース取得スクリプトのフルパス
            fetch_script = os.path.join(PRO_DB_DIR, 'fetch_industry_news.py')
            print(f"[INFO] ニュース取得開始: {fetch_script}")
            subprocess.run([sys.executable, fetch_script], check=True, capture_output=True)
            
            # 2. 要約生成スクリプトのフルパス
            gen_script = os.path.join(PRO_DB_DIR, 'generate_news_briefing.py')
            cmd = [sys.executable, gen_script]
            if api_key:
                cmd.extend(['--api-key', api_key])
            
            print(f"[INFO] 要約生成開始: {gen_script}")
            subprocess.run(cmd, check=True, capture_output=True)
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"status": "success"}).encode('utf-8'))
        except Exception as e:
            print(f"[ERROR] ニュース更新失敗: {e}")
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode('utf-8'))



if __name__ == '__main__':
    PORT = 8000
    server_address = ('', PORT)
    httpd = HTTPServer(server_address, UltimateAPIHandler)
    print(f"[SYSTEM] AGENS Ultimate API Server started on port {PORT}")
    httpd.serve_forever()
