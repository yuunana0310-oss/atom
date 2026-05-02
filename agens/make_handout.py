# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_red_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('  ' + text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(192, 0, 0)
    return p

# タイトル
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('新人職員研修　配布資料')
r.bold = True
r.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('移乗・移動介助の基本　〜安全・安楽・自立支援を実践する〜')
r.bold = True
r.font.size = Pt(12)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('対象：看護職員・介護職員　　所要時間：約30分　　医療法人 研修・教育委員会')
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# 研修目標
doc.add_heading('■ 本日の研修目標', level=2)
goals = [
    '① 移乗・移動介助の基本原則とボディメカニクスを理解する',
    '② 車椅子・ベッド間の移乗手順を正しく実施できる',
    '③ 歩行介助・立ち上がり介助を安全に行える',
    '④ 転倒・転落リスクを予測し予防策を講じられる',
    '⑤ 患者の尊厳と残存能力を尊重した介助ができる',
]
for g in goals:
    add_bullet(doc, g)
add_red_note(doc, '★ 患者の安全・安楽・自立支援が介助の3原則です')

# 介助の3原則
doc.add_heading('■ 介助の3原則', level=2)
table = doc.add_table(rows=2, cols=3)
table.style = 'Table Grid'
headers = ['安 全', '安 楽', '自立支援']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)

contents = [
    '・転倒・転落・皮膚損傷を防ぐ\n・環境整備を必ず行う\n・2人介助の基準を遵守する',
    '・患者の苦痛・不安を最小化\n・声かけ・同意を必ず得る\n・プライバシーを守る',
    '・できることは自分で行ってもらう\n・残存能力を活かす介助\n・ADL維持・向上を目指す',
]
for i, text in enumerate(contents):
    cell = table.rows[1].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

add_red_note(doc, '★ 声かけ・同意・プライバシー保護は必ず実施！')

# ボディメカニクス
doc.add_heading('■ ボディメカニクスの基本（腰痛予防・安全介助）', level=2)
bm = [
    ('① 支持基底面を広げる', '足を肩幅に開き、安定した姿勢を確保する'),
    ('② 重心を低くする', '膝を曲げて腰を落とし、重心を下げる'),
    ('③ 重心を近づける', '患者との距離を縮め、密着して介助する'),
    ('④ 大きな筋群を使う', '腰だけでなく脚・体幹の大きな筋肉を使う'),
    ('⑤ 身体をひとつにまとめる', '患者の手足を体に引き付けてから動かす'),
    ('⑥ ねじり動作を避ける', '足を向けてから体を回旋させる'),
]
for title_text, desc in bm:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(title_text + '　')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

# 介助前の確認事項
doc.add_heading('■ 介助前の確認事項', level=2)
table2 = doc.add_table(rows=2, cols=2)
table2.style = 'Table Grid'
for i, h in enumerate(['患者確認', '環境整備']):
    cell = table2.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10.5)

patient_items = '・患者の理解・同意を得る\n・バイタルサイン・体調確認\n・点滴・ドレーン・カテーテル確認\n・麻痺・疼痛・可動域の確認\n・体重・介助レベル確認'
env_items = '・ベッドの高さ調整\n・ブレーキ・ストッパーの確認\n・床の濡れ・障害物の除去\n・スリッパ→滑り止め靴に変更\n・プライバシーカーテンを閉める'
for i, text in enumerate([patient_items, env_items]):
    cell = table2.rows[1].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

add_red_note(doc, '⚠ 2人介助基準：全介助・体重60kg以上・不安定な患者は必ず2名で対応')
p = doc.add_paragraph()
r = p.add_run('  基本手順：確認 → 声かけ → 環境整備 → 介助 → 観察 → 記録')
r.font.size = Pt(10)

# 立ち上がり介助
doc.add_heading('■ 立ち上がり介助（ベッド端座位→立位）', level=2)
steps = [
    ('STEP 1', '端座位を確保する：ベッド端に座り、足底を床につける'),
    ('STEP 2', '前傾姿勢をとる：体を前に傾け重心を前方に移動、介助者は腰部を支持'),
    ('STEP 3', '立ち上がりを誘導：「せーの」の声かけで一緒に立ち上がる'),
    ('STEP 4', '立位の安定確認：めまい・ふらつきがないか確認、必要なら手すり・歩行器を使用'),
]
for step, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(step + '　')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

# 移乗手順
doc.add_heading('■ ベッド⇔車椅子 移乗手順（片麻痺患者）', level=2)
table3 = doc.add_table(rows=2, cols=2)
table3.style = 'Table Grid'
for i, h in enumerate(['ベッド → 車椅子', '車椅子 → ベッド']):
    cell = table3.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10.5)

bw_items = '1. 車椅子を健側に約30°の角度で設置\n2. ブレーキ確認・フットレストを上げる\n3. 患者を端座位にし足底を床につける\n4. 前傾姿勢→健側の手で介助しながら立上がり\n5. ゆっくり方向転換して車椅子に着座\n6. フットレストを下げ姿勢を整える'
wb_items = '（逆の手順で実施）\n1. 健側に30°でベッドに対して配置\n2. ブレーキ・フットレスト確認\n3. 前傾姿勢から立ち上がりを誘導\n4. 方向転換（介助者は患側に位置）\n5. ベッドに手をつかせてゆっくり着座\n6. 体位を整え座位バランスを確認'
for i, text in enumerate([bw_items, wb_items]):
    cell = table3.rows[1].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

# 歩行介助
doc.add_heading('■ 歩行介助の基本', level=2)
walk_points = [
    '介助者の位置：患側（麻痺のある側）の斜め後方',
    '片手で患者の腰または移乗ベルトを握る',
    '歩幅・速度は患者に合わせる',
    '出発前：立位保持→めまい確認→補助具用意',
    '歩行中：声かけしながら一定ペース・障害物注意',
    '方向転換：健側に向かってゆっくり回転（急な向き変えは禁止）',
    '着座時：椅子/ベッドに確実に触れてからゆっくり着座',
]
for item in walk_points:
    add_bullet(doc, item)

# 全介助（2名）
doc.add_heading('■ 全介助の移乗（2名介助）', level=2)
table4 = doc.add_table(rows=2, cols=2)
table4.style = 'Table Grid'
for i, h in enumerate(['リーダー役（頭側）', 'サポート役（足側）']):
    cell = table4.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10.5)

leader = '・上半身・頭頸部を支持\n・声かけ・タイミングの指示を出す\n・「いちにのさん」で全員同時に動く\n・気道・カテーテル類の安全を確認'
support = '・下肢・臀部を支持\n・リーダーの指示に従う\n・足先から骨盤まで確実に保持\n・移動方向を確認してから動く'
for i, text in enumerate([leader, support]):
    cell = table4.rows[1].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)

p = doc.add_paragraph()
r = p.add_run('  適応：全介助患者／体重60kg以上／不安定な患者／初回移乗')
r.font.size = Pt(10)

# 転倒・転落予防
doc.add_heading('■ 転倒・転落予防', level=2)
table5 = doc.add_table(rows=2, cols=3)
table5.style = 'Table Grid'
for i, h in enumerate(['患者側リスク', '環境側リスク', '予防策']):
    cell = table5.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)

r1 = '・筋力低下・バランス障害\n・認知症・見当識障害\n・薬剤影響（睡眠薬等）\n・起立性低血圧\n・排泄への切迫感'
r2 = '・床の濡れ・障害物\n・不適切な照明\n・ベッド高さ不適切\n・ナースコール未設置\n・滑りやすい履物'
r3 = '・転倒リスクスコア評価（毎日）\n・離床センサー設置\n・低床ベッドの使用\n・ナースコールを手の届く位置へ\n・排泄パターン把握しラウンド強化'
for i, text in enumerate([r1, r2, r3]):
    cell = table5.rows[1].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9)

add_red_note(doc, '転倒・転落発生時：患者の安全確保 → ナースコール → 状態確認 → 報告 → 記録')

# よくあるミス
doc.add_heading('■ よくあるミスと対策', level=2)
table6 = doc.add_table(rows=1, cols=3)
table6.style = 'Table Grid'
for i, h in enumerate(['よくあるミス', '起こりうる問題', '対策']):
    cell = table6.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)

mistakes = [
    ('× 声かけを忘れる', '患者の心理的不安・筋緊張が増す', '○ 必ず「○○しますね」と伝えてから開始'),
    ('× ブレーキを忘れる', '移乗中に車椅子が動いて転落', '○ 移乗前の確認を手順書通りに実施'),
    ('× 腰だけで持ち上げる', '介助者の腰痛・ぎっくり腰', '○ ボディメカニクスを意識。足・体幹を使う'),
    ('× 急ぎすぎる', '患者が不安・バランス崩す', '○ 患者のペースに合わせ、ゆっくり確実に'),
    ('× 1人で無理をする', '患者・介助者ともに受傷リスク', '○ 迷ったら必ず同僚に声をかける'),
    ('× カテーテルを忘れる', 'ドレーン抜去・皮膚損傷', '○ 移乗前にルート類の長さ・固定を確認'),
]
for m in mistakes:
    row = table6.add_row()
    for i, text in enumerate(m):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        if i == 2:
            run.font.color.rgb = RGBColor(0, 112, 0)

# 確認テスト
doc.add_heading('■ 確認テスト', level=2)
qs = [
    ('Q1', 'ボディメカニクスの目的は何ですか？', '  ①患者の安楽　②介助者の腰痛予防　③両方　④どちらでもない'),
    ('Q2', '片麻痺患者への車椅子移乗で、車椅子を置く位置はどちら側？', '  ①患側（麻痺側）　②健側（麻痺のない側）　③どちらでも同じ'),
    ('Q3', '2名介助が必要な条件を1つ答えてください。', '（                                                              ）'),
    ('Q4', '移乗介助前に必ず確認すべき3つのことを答えてください。', '（１）                  （２）                  （３）                  '),
    ('Q5', '転倒が発生した際の対応の優先順位を答えてください。', '（                                                              ）'),
]
for q_num, q_text, options in qs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(q_num + '　')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(q_text)
    run2.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run3 = p2.add_run(options)
    run3.font.size = Pt(9.5)
    run3.font.color.rgb = RGBColor(80, 80, 80)

# まとめ
doc.add_heading('■ 本日のまとめ', level=2)
summary = [
    ('1', '介助の3原則', '安全・安楽・自立支援を常に意識する'),
    ('2', 'ボディメカニクス', '正しい姿勢で患者も介助者も守る'),
    ('3', '確認・声かけ', '介助前の確認と患者への声かけを徹底する'),
    ('4', '移乗手順の遵守', '車椅子の角度・ブレーキ・フットレストを確認'),
    ('5', '2名介助の判断', '一人で無理せず、迷ったら必ず声をかける'),
    ('6', '記録・報告', '気づいたことはすぐに記録・報告・連絡'),
]
for num, key, value in summary:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(num + '　' + key + '：')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(value)
    run2.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('患者の「安全・安楽・自立支援」のために、今日から実践しましょう！')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.save('e:/新人/移乗移動介助_配布資料.docx')
print('配布資料を保存しました: e:/新人/移乗移動介助_配布資料.docx')
