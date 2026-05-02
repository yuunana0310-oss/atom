# -*- coding: utf-8 -*-
"""
発表原稿.docx に新規追加セクションの発表原稿を挿入する。
  1. PART 01内 — 制服SNS投稿禁止の説明（既存のSNS段落の直後）
  2. PART 02とPART 03の間 — パワーハラスメント新セクション

書式は既存の 発表原稿.docx と完全に一致させる。
"""
import docx, io, sys, copy
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc = docx.Document('発表原稿.docx')

# ──────────────────────────────────────────────────────────
# Helper functions — exact match to existing format
# ──────────────────────────────────────────────────────────

def new_p():
    return OxmlElement('w:p')

def set_spacing(p, before_pt=None, after_pt=None):
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if before_pt is not None:
        sp.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt is not None:
        sp.set(qn('w:after'), str(int(after_pt * 20)))

def set_indent(p, left_pt):
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), str(int(left_pt * 20)))

def set_shading(p, fill_hex):
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def add_run(p, text, bold=False, size_pt=None, color_hex=None):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if size_pt is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        rPr.append(szCs)
    if color_hex:
        col = OxmlElement('w:color')
        col.set(qn('w:val'), color_hex)
        rPr.append(col)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)

# ──────────────────────────────────────────────────────────
# Paragraph builders matching existing styles exactly
# ──────────────────────────────────────────────────────────

def make_section_header(text):
    """■ PARTヘッダー: shd=333333, bold 12pt white, sp_after=6pt"""
    p = new_p()
    set_spacing(p, after_pt=6)
    set_shading(p, '333333')
    add_run(p, text, bold=True, size_pt=12, color_hex='FFFFFF')
    return p

def make_page_ref(text):
    """【資料○ページ】: shd=F0F0F0, 9pt color=555555, ind=14.2pt"""
    p = new_p()
    set_indent(p, left_pt=14.17)   # 179705 EMU / 12700 = 14.15pt
    set_shading(p, 'F0F0F0')
    add_run(p, text, size_pt=9, color_hex='555555')
    return p

def make_body(text, bold=False, indent=False):
    """本文: 11pt, sp_before=1pt sp_after=3pt, indent=14.2pt (optional)"""
    p = new_p()
    set_spacing(p, before_pt=1, after_pt=3)
    if indent:
        set_indent(p, left_pt=14.17)
    add_run(p, text, bold=bold, size_pt=11)
    return p

def make_blank():
    """空行"""
    return new_p()

# ──────────────────────────────────────────────────────────
# PART 1: SNS制服投稿 — PART 01内、既存SNS段落の直後に挿入
# ──────────────────────────────────────────────────────────
# 挿入位置: 「SNSへの投稿も同様です。〜は通用しません。」の直後（廊下の段落の前）

ref_sns = None
for para in doc.paragraphs:
    if 'SNSへの投稿も同様' in para.text:
        ref_sns = para._p
        break
assert ref_sns is not None, 'SNS paragraph not found'

sns_paras = [
    make_body('★ 今回、新たに確認してほしいルールがあります。制服姿や院内の写真・動画をプライベートSNSに投稿することも禁止です。', bold=True),
    make_body('「自分が映っているだけだから問題ない」——そう思う方もいるかもしれません。でも、それは誤解です。'),
    make_body('制服や背景から勤務先が特定される。写り込んだ書類や画面が患者情報の漏洩になる。', indent=True),
    make_body('ご家族の姿や病院の設備が映り込み、意図せず患者情報を公開してしまうケースもあります。', indent=True),
    make_body('さらに——位置情報や行動パターンから、スタッフ自身がストーカー被害に遭う事例も実際に起きています。', indent=True),
    make_body('★ 制服・名札・院内設備が映り込む写真・動画は、プライベートSNSへの投稿を一切禁止します。今日から必ず守ってください。', bold=True),
]

# insert in order after ref_sns
prev = ref_sns
for p in sns_paras:
    prev.addnext(p)
    prev = p

print(f'Inserted {len(sns_paras)} SNS paragraphs after PART 01 SNS line')

# ──────────────────────────────────────────────────────────
# PART 2: パワーハラスメント — PART 02 と PART 03 の間に挿入
# ──────────────────────────────────────────────────────────
# 挿入位置: 「★ もう一度だけ言います。カスハラは〜」の直後（空行＋PART 03ヘッダーの前）

ref_part03 = None
for para in doc.paragraphs:
    if 'PART 03' in para.text and '法令遵守' in para.text:
        ref_part03 = para._p
        break
assert ref_part03 is not None, 'PART 03 not found'

pw_paras = [
    make_blank(),
    make_section_header('■  PART 02-2　パワーハラスメント（パワハラ）（約4分）'),
    make_page_ref('【資料5-2ページ】'),
    make_blank(),
    make_body('次にパワーハラスメントについて確認します。資料の02-2ページを開いてください。'),
    make_body('パワハラとはどういうものか。厚生労働省の定義では、3つの要件をすべて満たす言動がパワハラとされています。'),
    make_body('★「①優越的な関係を背景とした言動」「②業務上必要かつ相当な範囲を超えたもの」「③労働者の就業環境が害されるもの」。', bold=True),
    make_body('この3つが揃ったとき、パワハラと判断されます。'),
    make_blank(),
    make_body('パワハラには6つの類型があります。'),
    make_body('身体的な攻撃——暴行・傷害。これは論外ですが、「軽く肩を叩く」程度でも繰り返せば問題になります。', indent=True),
    make_body('精神的な攻撃——脅迫・侮辱・暴言。「なんでこんなこともできないんだ」という怒鳴り声は、指導ではなくパワハラです。', indent=True),
    make_body('人間関係からの切り離し——無視・仲間外し・隔離。挨拶を無視し続けること、情報共有から外すことも該当します。', indent=True),
    make_body('過大な要求——明らかに達成不可能な業務を強制すること。', indent=True),
    make_body('過小な要求——能力・経験を無視した単純な仕事しか与えないこと。', indent=True),
    make_body('個の侵害——プライベートへの過度な立ち入り。休日の予定を執拗に聞く、SNSを監視するなどが当たります。', indent=True),
    make_blank(),
    make_body('★「指導のつもりだった」「本人のためを思っていた」では通りません。相手が精神的ダメージを受け、就業環境が害されていれば、パワハラと判断される可能性があります。', bold=True),
    make_blank(),
    make_body('では、パワハラをしないために何を意識すればよいか。'),
    make_body('まず——「人格」ではなく「行動」を指摘してください。「あなたはダメだ」はNG。「この書類の○○の部分を直してほしい」と、具体的な行動を指摘することが大切です。', indent=True),
    make_body('注意・指導は必ず個別・非公開の場で行うこと。人前での叱責は相手の尊厳を傷つけ、パワハラと受け取られます。', indent=True),
    make_body('「これくらい当たり前」という感覚を疑ってください。自分の「当たり前」が、相手には過大な負担である場合があります。', indent=True),
    make_body('★ 自分の言動を振り返る習慣を持つこと。そして迷ったら——一人で判断せず、相談窓口に声をかけてください。', bold=True, indent=True),
    make_blank(),
    make_body('カスハラは「外からくるハラスメント」、パワハラは「内からくるハラスメント」です。'),
    make_body('どちらも、職場全体が安心して働ける環境を守るために、一人ひとりの意識が必要です。'),
]

# insert in order before PART 03 header
# addprevious(p) puts p immediately before ref_part03 each time.
# Forward order (p0 first) → p0 pushed furthest, pN stays nearest → reading: p0,p1,...,pN,ref ✓
for p in pw_paras:
    ref_part03.addprevious(p)

print(f'Inserted {len(pw_paras)} パワハラ paragraphs before PART 03')

doc.save('発表原稿_new.docx')
print('Saved: 発表原稿_new.docx')
