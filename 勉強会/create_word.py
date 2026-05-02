from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ===== ユーティリティ =====

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph_border_bottom(para, color='000000', sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_shading(para, fill_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)

def set_para_border(para, color='000000', sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

def set_para_left_border(para, color='000000', sz=24):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def set_run_font(run, font_name='Meiryo'):
    run.font.name = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def heading_para(doc, text, level=1, font_size=16, bold=True, color='111111', border_bottom=True, space_before=200, space_after=100):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before / 20)
    p.paragraph_format.space_after = Pt(space_after / 20)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)
    if border_bottom:
        add_paragraph_border_bottom(p, sz=18)
    return p

def body_para(doc, text, font_size=10.5, bold=False, color='111111', indent=None, space_before=40, space_after=40):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before / 20)
    p.paragraph_format.space_after = Pt(space_after / 20)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)
    return p, run

def bullet_para(doc, text, bullet='▶', font_size=10, indent_cm=0.5, bold_part=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm + 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run_bullet = p.add_run(f'{bullet}  ')
    run_bullet.font.size = Pt(font_size - 1)
    set_run_font(run_bullet)
    run_text = p.add_run(text)
    run_text.font.size = Pt(font_size)
    if bold_part:
        run_text.bold = True
    set_run_font(run_text)
    return p

def ng_bullet(doc, text, font_size=10, indent_cm=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm + 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run_x = p.add_run('✕  ')
    run_x.font.size = Pt(font_size)
    run_x.bold = True
    set_run_font(run_x)
    run_text = p.add_run(text)
    run_text.font.size = Pt(font_size)
    set_run_font(run_text)
    return p

def step_bullet(doc, num, text, font_size=10, indent_cm=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm + 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_n = p.add_run(f'【{num}】  ')
    run_n.font.size = Pt(font_size)
    run_n.bold = True
    set_run_font(run_n)
    run_text = p.add_run(text)
    run_text.font.size = Pt(font_size)
    set_run_font(run_text)
    return p

def dark_box(doc, title, body_lines, font_size=10):
    """黒背景ボックス（パラグラフ単位で再現）"""
    lines = [title] + body_lines
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0) if i > 0 else Pt(4)
        p.paragraph_format.space_after = Pt(0) if i < len(lines)-1 else Pt(6)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.4)
        set_para_shading(p, '111111')
        run = p.add_run(line)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(font_size)
        if i == 0:
            run.bold = True
        set_run_font(run)

def bordered_box(doc, label, body_lines, font_size=10):
    """枠線ボックス"""
    for i, line in enumerate(body_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0) if i > 0 else Pt(4)
        p.paragraph_format.space_after = Pt(0) if i < len(body_lines)-1 else Pt(6)
        p.paragraph_format.left_indent = Cm(0.4)
        set_para_border(p)
        if i == 0 and label:
            run_lbl = p.add_run(f'【{label}】  ')
            run_lbl.font.size = Pt(font_size - 0.5)
            run_lbl.bold = True
            set_run_font(run_lbl)
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        set_run_font(run)

def left_border_box(doc, lines, font_size=10):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        set_para_left_border(p)
        set_para_shading(p, 'F5F5F5')
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        set_run_font(run)

def sub_heading(doc, text, font_size=11, space_before=160):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before / 20)
    p.paragraph_format.space_after = Pt(4)
    add_paragraph_border_bottom(p, sz=8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    set_run_font(run)
    return p

def page_break(doc):
    doc.add_page_break()

def spacer(doc, size=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('')
    run.font.size = Pt(size)

# ===== ドキュメント作成 =====

def create_handout():
    doc = Document()

    # ページ余白
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ===== 表紙 =====
    spacer(doc, 40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026年度　院内必修研修')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_run_font(run)

    spacer(doc, 30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ['個人情報保護', 'ハラスメント予防', '法令遵守']:
        run = p.add_run(line + '\n')
        run.font.size = Pt(24)
        run.bold = True
        set_run_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('すべての病院スタッフへ')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    set_run_font(run)

    spacer(doc, 20)
    # 区切り線
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border_bottom(p, sz=18)
    run = p.add_run('')
    run.font.size = Pt(4)

    spacer(doc, 10)
    # テーマバッジ（表形式）
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, badge in enumerate(['📋  個人情報保護', '🛡  ハラスメント予防', '⚖  法令遵守']):
        cell = tbl.cell(0, i)
        cell.width = Cm(5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(badge)
        run.bold = True
        run.font.size = Pt(10)
        set_run_font(run)
        border_val = {'val': 'single', 'sz': 12, 'color': '111111'}
        set_cell_border(cell, top=border_val, bottom=border_val, left=border_val, right=border_val)

    spacer(doc, 30)
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_meta.add_run('研修時間：約30分　／　配布資料あり\n対象：全職員　／　実施部署：________________')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_run_font(run)

    # ナノバナナ指示（表紙）
    spacer(doc, 20)
    bordered_box(doc, 'ナノバナナ イラスト指示（表紙）',
        ['「鍵・盾・法律の本を持った3人のシルエット人物」',
         'モノクロ線画・横並び・横150px×縦80px程度',
         '→ cover_illust.png を作成してここに挿入してください'])

    page_break(doc)

    # ===== アジェンダ =====
    heading_para(doc, '00　本日のアジェンダ', font_size=15)
    body_para(doc, '患者さんとスタッフを守るために——今日の研修で確認すべきことを整理しました。\n配布資料と合わせて、ご自身の業務を振り返りながらご参加ください。',
              font_size=10, color='444444')
    spacer(doc, 6)

    # アジェンダ表
    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ('PART 01', '個人情報保護\n患者情報の正しい扱い方・漏洩防止'),
        ('PART 02', 'ハラスメント予防\n職場・対患者 ／ カスハラ重点解説'),
        ('PART 03', '法令遵守\n医療現場に関わる主要法令と義務'),
        ('WRAP UP', 'セルフチェックリスト\n自己確認チェックリスト'),
    ]
    for r, (label, desc) in enumerate(items):
        c0 = tbl2.cell(r, 0)
        c1 = tbl2.cell(r, 1)
        c0.width = Cm(3)
        c1.width = Cm(12)
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(label)
        run0.bold = True
        run0.font.size = Pt(9)
        set_run_font(run0)
        set_cell_bg(c0, 'F0F0F0')
        p1 = c1.paragraphs[0]
        for part in desc.split('\n'):
            run1 = p1.add_run(part)
            run1.font.size = Pt(10)
            set_run_font(run1)
            if part != desc.split('\n')[-1]:
                p1.add_run('\n')

    spacer(doc, 8)
    dark_box(doc, '研修のゴール', [
        '① 患者情報を「なぜ・どう守るか」が言える',
        '② ハラスメント（特にカスハラ）の定義と対応を知っている',
        '③ 違反した場合のリスクを理解し、迷ったとき相談できる',
    ])

    page_break(doc)

    # ===== 個人情報保護 =====
    heading_para(doc, '01　個人情報保護', font_size=15)

    # ナノバナナ指示
    bordered_box(doc, 'ナノバナナ指示（個人情報ページ）',
        ['「鍵のかかったファイルを抱える人物」シルエット線画',
         '横80px×縦80px・右寄せ → illust_privacy.png を挿入'])

    sub_heading(doc, '「個人情報」とは何か？')

    # 2列レイアウト（表で代用）
    tbl3 = doc.add_table(rows=1, cols=2)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left = tbl3.cell(0, 0)
    c_right = tbl3.cell(0, 1)
    c_left.width = Cm(7.5)
    c_right.width = Cm(7.5)

    for cell, title, items in [
        (c_left, '個人情報の定義', ['氏名・生年月日・住所・電話番号', '診断名・病歴・検査結果・投薬内容', '顔写真・音声・映像', 'マイナンバー・保険証番号']),
        (c_right, '「要配慮個人情報」（特に厳格）', ['傷病・障害・健康診断結果', '精神疾患・感染症の罹患情報', '手術・検査・処方の記録', '⇒ 特に厳格な取り扱いが必要']),
    ]:
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        set_run_font(run)
        for item in items:
            p2 = cell.add_paragraph()
            r1 = p2.add_run('▶  ')
            r1.font.size = Pt(8)
            set_run_font(r1)
            r2 = p2.add_run(item)
            r2.font.size = Pt(9.5)
            set_run_font(r2)
            p2.paragraph_format.left_indent = Cm(0.2)

    spacer(doc, 8)
    sub_heading(doc, '患者情報取り扱い 5原則')

    # 5原則（横並び表）
    tbl4 = doc.add_table(rows=2, cols=5)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    nums = ['1', '2', '3', '4', '5']
    labels = ['目的外\n使用禁止', '最小限の\nアクセス', '第三者提供\nの制限', '安全管理\n措置', '漏洩時の\n即時報告']
    for i in range(5):
        cn = tbl4.cell(0, i)
        cl = tbl4.cell(1, i)
        cn.width = Cm(3)
        cl.width = Cm(3)
        pn = cn.paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = pn.add_run(nums[i])
        rn.bold = True
        rn.font.size = Pt(18)
        set_run_font(rn)
        set_cell_bg(cn, 'F0F0F0')
        pl = cl.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(labels[i])
        rl.bold = True
        rl.font.size = Pt(9)
        set_run_font(rl)

    spacer(doc, 8)
    sub_heading(doc, 'やってはいけないこと NG10')

    ng_items = [
        '院外でのカルテ・書類の閲覧・持ち出し',
        '私用スマホで患者情報を撮影・保存',
        'SNSへの患者関連投稿（特定不能でも不可）',
        '廊下・エレベーターでの患者情報の会話',
        '無関係な患者のカルテを閲覧する',
        '書類の放置・施錠忘れ',
        '患者情報を含むメールを誤送信',
        '家族・知人への患者情報の口頭伝達',
        '退職後の情報の持ち出し・利用',
        '業務用PCの共有パスワード使用',
    ]
    for item in ng_items:
        ng_bullet(doc, item)

    spacer(doc, 6)
    dark_box(doc, '違反した場合のリスク', [
        '個人情報保護法違反 → 刑事罰（1年以下の懲役または50万円以下の罰金）・行政指導・民事賠償責任',
        '当院への社会的信用失墜・患者との信頼関係の崩壊',
    ])

    page_break(doc)

    # ===== ハラスメント予防 =====
    heading_para(doc, '02　ハラスメント予防', font_size=15)
    sub_heading(doc, 'ハラスメントの主な種類')

    # ハラスメント表
    tbl5 = doc.add_table(rows=5, cols=3)
    tbl5.style = 'Table Grid'
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['種類', '定義', '病院内の例']
    for i, h in enumerate(headers):
        cell = tbl5.cell(0, i)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        set_run_font(run)
        set_cell_bg(cell, '111111')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ('パワハラ', '優越的立場を利用した言動で就業環境を害する', '「なんでこんなこともできないんだ」と怒鳴る'),
        ('セクハラ', '性的言動により就業環境を害する', '体型や外見についての不適切な発言'),
        ('マタハラ', '妊娠・出産・育休に関する不利益な取り扱い', '「妊娠するタイミングが悪い」などの発言'),
        ('カスハラ ★重点', '患者・家族等による著しく不当な要求・言動', '長時間の怒鳴り・土下座要求・SNS投稿脅迫'),
    ]
    for r, (t, d, e) in enumerate(rows_data):
        for c_idx, text in enumerate([t, d, e]):
            cell = tbl5.cell(r + 1, c_idx)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            if c_idx == 0:
                run.bold = True
            if r == 3 and c_idx == 0:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            set_run_font(run)
        if r == 3:
            set_cell_bg(tbl5.cell(r + 1, 0), 'E0E0E0')

    page_break(doc)

    # ===== カスハラ重点 =====
    heading_para(doc, '02　カスタマーハラスメント（カスハラ）重点解説', font_size=14)

    bordered_box(doc, 'ナノバナナ指示（カスハラページ）',
        ['「受付カウンターで怒っている人物と盾を持つスタッフ」シルエット線画・横長',
         '横200px×縦70px → illust_kushara.png を挿入'])

    spacer(doc, 4)
    bordered_box(doc, '定義',
        ['顧客・患者・その家族等が行う、社会通念上不相当な言動で、',
         '労働者の就業環境が害されるもの',
         '（厚生労働省 2024年指針より）'])

    spacer(doc, 6)
    sub_heading(doc, '病院で起こりやすいカスハラの具体例')

    tbl6 = doc.add_table(rows=2, cols=2)
    tbl6.style = 'Table Grid'
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    examples = [
        ('① 暴言・脅迫系', [
            '「訴えるぞ」「SNSで晒す」などの脅し',
            '大声での怒鳴り・罵倒・人格否定',
            '「上を呼べ」の連続・執拗なクレーム',
        ]),
        ('② 不当要求系', [
            '土下座・謝罪文の強要',
            '医学的に不必要な処置・薬の要求',
            '特別扱い（優先診察など）の強要',
        ]),
        ('③ 身体的行為系', [
            '物を投げる・叩くなどの暴力行為',
            'スタッフへの不審な接触・追跡',
        ]),
        ('④ 長時間拘束系', [
            '同じ要求を繰り返し何時間も居座る',
            '電話での長時間クレーム',
        ]),
    ]
    positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for (r, c), (title, items) in zip(positions, examples):
        cell = tbl6.cell(r, c)
        cell.width = Cm(7.5)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        set_run_font(run)
        for item in items:
            p2 = cell.add_paragraph()
            r1 = p2.add_run('✕  ')
            r1.font.size = Pt(9)
            r1.bold = True
            set_run_font(r1)
            r2 = p2.add_run(item)
            r2.font.size = Pt(9.5)
            set_run_font(r2)
            p2.paragraph_format.left_indent = Cm(0.2)

    spacer(doc, 8)
    sub_heading(doc, 'カスハラへの対応フロー（スタッフ向け）')

    steps = [
        ('STEP 1', '不当な言動に気づく — 「これはクレームか、カスハラか」を判断する'),
        ('STEP 2', '一人で抱え込まない — すぐに同僚・先輩・リーダーへ声をかける'),
        ('STEP 3', '複数人で対応 — 必ず2名以上。単独対応しない'),
        ('STEP 4', '事実を記録する — 日時・言動の内容を具体的にメモ'),
        ('STEP 5', '上長・相談窓口へ報告 — 組織として対応する'),
    ]
    for i, (label, desc) in enumerate(steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if i == len(steps) - 1:
            set_para_shading(p, '111111')
            r1 = p.add_run(f'{label}　')
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_run_font(r1)
            r2 = p.add_run(desc)
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_run_font(r2)
        else:
            set_para_border(p, sz=6)
            r1 = p.add_run(f'{label}　')
            r1.bold = True
            r1.font.size = Pt(10)
            set_run_font(r1)
            r2 = p.add_run(desc)
            r2.font.size = Pt(10)
            set_run_font(r2)
        if i < len(steps) - 1:
            p_arrow = doc.add_paragraph()
            p_arrow.paragraph_format.left_indent = Cm(0.5)
            p_arrow.paragraph_format.space_before = Pt(0)
            p_arrow.paragraph_format.space_after = Pt(0)
            run_a = p_arrow.add_run('　　　　▼')
            run_a.font.size = Pt(9)
            run_a.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            set_run_font(run_a)

    spacer(doc, 6)
    dark_box(doc, '組織の責任：使用者はスタッフを守る義務がある', [
        '2024年労働施策総合推進法の指針改正により、事業者はカスハラへの対策・相談体制・被害スタッフへのフォローが義務化。',
        'スタッフ個人が我慢する問題ではなく、組織として取り組むべき課題です。',
    ])

    spacer(doc, 4)
    bordered_box(doc, '覚えておこう',
        ['「正当なクレーム」と「カスハラ」は違います。',
         '患者の不満・要望は丁寧に対応する。しかし、社会通念を超えた言動は受け入れないのが正しい姿勢です。'])

    page_break(doc)

    # ===== 法令遵守 =====
    heading_para(doc, '03　法令遵守（コンプライアンス）', font_size=15)

    bordered_box(doc, 'ナノバナナ指示（法令遵守ページ）',
        ['「天秤の左に病院、右に法律書が乗ったシンプル線画」モノクロ',
         '横90px×縦80px → illust_compliance.png を挿入'])

    spacer(doc, 4)
    sub_heading(doc, '医療現場に関わる主要法令')

    tbl7 = doc.add_table(rows=6, cols=2)
    tbl7.style = 'Table Grid'
    tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['法令', 'スタッフが知るべきポイント']):
        cell = tbl7.cell(0, i)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_run_font(run)
        set_cell_bg(cell, '111111')
    tbl7.cell(0, 0).width = Cm(5)
    tbl7.cell(0, 1).width = Cm(10)

    laws = [
        ('個人情報保護法', '患者情報の目的外利用・第三者提供の禁止。漏洩時は72時間以内に報告義務'),
        ('医療法', '医療記録の正確な保管義務・インフォームドコンセントの確保'),
        ('労働施策総合推進法\n（パワハラ防止法）', '職場のパワハラ・カスハラ対策が事業者の義務。相談窓口の設置が必須'),
        ('医師法・保健師助産師看護師法', '守秘義務（業務で知った情報を漏らしてはならない）。退職後も継続'),
        ('不正競争防止法', '院内の業務秘密の持ち出し・漏洩の禁止'),
    ]
    for r, (law, point) in enumerate(laws):
        c0 = tbl7.cell(r + 1, 0)
        c1 = tbl7.cell(r + 1, 1)
        c0.width = Cm(5)
        c1.width = Cm(10)
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(law)
        run0.bold = True
        run0.font.size = Pt(9)
        set_run_font(run0)
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(point)
        run1.font.size = Pt(9)
        set_run_font(run1)
        if r % 2 == 1:
            set_cell_bg(c0, 'F5F5F5')
            set_cell_bg(c1, 'F5F5F5')

    spacer(doc, 8)
    sub_heading(doc, '違反した場合に起こること')

    tbl8 = doc.add_table(rows=1, cols=3)
    tbl8.style = 'Table Grid'
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    penalties = [
        ('刑事罰', '懲役・罰金。前科がつく場合も。'),
        ('行政処分', '免許停止・取り消し・業務停止。'),
        ('民事賠償', '患者・院への損害賠償請求。'),
    ]
    for i, (title, desc) in enumerate(penalties):
        cell = tbl8.cell(0, i)
        cell.width = Cm(5)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_run_font(run)
        set_cell_bg(cell, '111111')
        p2 = cell.add_paragraph()
        run2 = p2.add_run(desc)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_run_font(run2)

    spacer(doc, 8)
    sub_heading(doc, '迷ったとき・気づいたときは「報告・相談」が最優先')

    for num, text in [
        ('1', '「これはまずいかも？」と思ったら、まず直属の上長へ報告する'),
        ('2', '上長への相談が難しい場合は院内コンプライアンス相談窓口を利用する'),
        ('3', '隠さない・後回しにしない —— 早期対応が被害を最小化する'),
    ]:
        step_bullet(doc, num, text)

    spacer(doc, 6)
    bordered_box(doc, '院内相談窓口',
        ['個人情報・コンプライアンス：　担当部署 ___________　内線 ___________',
         'ハラスメント：　担当部署 ___________　内線 ___________'])

    page_break(doc)

    # ===== セルフチェックリスト =====
    heading_para(doc, '✓　今日から使えるセルフチェックリスト', font_size=14)
    body_para(doc, '研修後、自分の行動を振り返り □ にチェックを入れてください。チェックできない項目が改善ポイントです。',
              font_size=10, color='444444')
    spacer(doc, 6)

    sub_heading(doc, '個人情報保護', font_size=10.5)
    for item in [
        '患者情報を院外・廊下などで話していない',
        '私用スマホで患者情報を撮影・保存していない',
        '業務終了時に書類を施錠保管している',
        'SNSへの患者関連投稿を一切していない',
        '担当患者以外のカルテは閲覧していない',
    ]:
        bullet_para(doc, item, bullet='□')

    sub_heading(doc, '法令遵守', font_size=10.5, space_before=120)
    for item in [
        '守秘義務（退職後も含む）を認識している',
        '業務上の書類・情報を無断で持ち出していない',
        '迷ったときは上長・相談窓口に報告できている',
        'インシデントは速やかに報告している',
    ]:
        bullet_para(doc, item, bullet='□')

    sub_heading(doc, 'ハラスメント予防', font_size=10.5, space_before=120)
    for item in [
        '指導の際に怒鳴る・人格否定をしていない',
        '同僚の妊娠・育休を否定的に発言していない',
        'スタッフ間で相手が嫌がる言動をしていない',
        '患者・家族から不当な言動を受けたとき報告している',
        'カスハラと正当なクレームの区別ができる',
        'カスハラ対応は複数人で行っている',
    ]:
        bullet_para(doc, item, bullet='□')

    spacer(doc, 8)
    dark_box(doc, 'まとめ：3つのメッセージ', [
        '① 患者情報はあなたが「預かっている」もの。慎重に扱ってください。',
        '② ハラスメントは職場の問題。一人で抱え込まず、組織で対応しましょう。',
        '③ 迷ったら報告。隠さないことが、自分と患者と病院を守ります。',
    ])

    spacer(doc, 8)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border_bottom(p_foot, sz=4)
    run = p_foot.add_run('院内研修資料 2026年度　／　無断複製・外部持ち出し禁止')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    set_run_font(run)

    doc.save('C:/Users/yuuna/agens/勉強会/配布資料.docx')
    print('配布資料.docx を作成しました')


def create_script():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # タイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('発表原稿')
    run.bold = True
    run.font.size = Pt(20)
    set_run_font(run)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('病院スタッフ必修研修　個人情報保護・ハラスメント予防・法令遵守')
    run2.font.size = Pt(12)
    set_run_font(run2)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('目安：約20分（早読み）')
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_run_font(run3)
    add_paragraph_border_bottom(p3)

    # 読み方メモ
    spacer(doc, 6)
    bordered_box(doc, '読み方メモ',
        ['【　】内は配布資料の参照ページ',
         '／ は少し間を置くポイント',
         '★ 記号がある箇所はやや強調してゆっくり読む',
         '全体で約5,800字　1分300字ペースで約19〜20分'])

    def script_section(title):
        spacer(doc, 6)
        p = doc.add_paragraph()
        run = p.add_run(f'■  {title}')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_run_font(run)
        set_para_shading(p, '333333')
        p.paragraph_format.space_after = Pt(6)

    def script_line(text, indent=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.font.size = Pt(11)
        set_run_font(run)

    def script_bold(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        set_run_font(run)

    def script_note(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        set_para_shading(p, 'F0F0F0')
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True
        set_run_font(run)

    # ===== イントロ =====
    script_section('開始・イントロ（約1分）')
    script_line('みなさん、お疲れのところお集まりいただきありがとうございます。')
    script_line('本日は約30分、研修の時間をいただきます。')
    spacer(doc, 2)
    script_line('テーマは「個人情報保護」「ハラスメント予防」「法令遵守」の3つです。')
    script_line('どれも「知っているつもり」になりやすいテーマですが、実際の現場で問題になるのは、「知らなかった」ではなく「まあいいか」と思った瞬間です。')
    spacer(doc, 2)
    script_line('今日の研修で確認したいのは3点だけです。資料の2ページ目をご覧ください。')
    spacer(doc, 2)
    script_line('患者情報を「なぜ・どう守るか」が自分の言葉で説明できること。', indent=True)
    script_line('ハラスメント、特にカスハラの定義と対応を知っていること。', indent=True)
    script_line('違反したときのリスクを理解して、迷ったとき相談できること。', indent=True)
    spacer(doc, 2)
    script_line('この3つだけ、今日持ち帰っていただければ十分です。では始めます。')

    # ===== 個人情報保護 =====
    script_section('PART 01　個人情報保護（約6分）')
    script_note('【資料3ページ】')
    spacer(doc, 2)
    script_line('まず「個人情報とは何か」というところから確認します。')
    script_line('個人情報とは、「特定の個人を識別できる情報」のことです。')
    script_line('氏名・生年月日・住所・電話番号、これはみなさんご存知だと思います。でも病院で扱う情報はそれだけじゃない。')
    spacer(doc, 2)
    script_line('診断名・病歴・検査結果・投薬内容、これらはすべて個人情報です。')
    script_bold('★ さらに病院が扱うのは、法律上「要配慮個人情報」と呼ばれるもので、一般の個人情報よりさらに厳格な取り扱いが求められます。')
    script_line('精神疾患の有無、感染症の罹患、手術記録——これらは漏れただけで患者さんの人生に深刻な影響を与えかねない情報です。')
    spacer(doc, 2)
    script_line('資料の「患者情報取り扱い5原則」を見てください。')
    spacer(doc, 2)
    script_line('原則の1番目は「目的外使用の禁止」。患者さんから得た情報は、その診療のためだけに使う。それが大前提です。', indent=True)
    script_line('2番目は「最小限のアクセス」。自分の担当でない患者さんのカルテを「ちょっと見てみようかな」は、それだけでアウトです。', indent=True)
    script_line('3番目は「第三者提供の制限」。家族からの問い合わせでも、本人の同意なく情報を渡すことは原則できません。', indent=True)
    script_line('4番目は「安全管理措置」。書類の放置、PCの施錠忘れ、これも違反になり得ます。', indent=True)
    script_line('5番目は「漏洩時の即時報告」。問題が起きたとき、隠すより早く報告したほうが被害を最小化できます。', indent=True)
    spacer(doc, 2)
    script_line('次に「やってはいけないこと」を確認します。')
    spacer(doc, 2)
    script_bold('★ 一番やりがちなのが、スマートフォンの問題です。')
    script_line('「ちょっとメモ代わりに」と患者さんの情報を私用スマホで撮影する、これは完全にアウトです。撮った瞬間に違反。')
    spacer(doc, 2)
    script_line('SNSへの投稿も同様です。「患者さんの名前は書いてないから大丈夫」は通用しません。')
    script_line('廊下やエレベーターでの会話も要注意。「あの患者さん、実はね」という会話を誰が聞いているかわかりません。')
    spacer(doc, 2)
    script_bold('★ 見落としがちなのが退職後の守秘義務です。')
    script_line('辞めたあとも、在職中に知った患者情報を話すことは禁じられています。これは一生続きます。')
    spacer(doc, 2)
    script_line('違反した場合ですが、個人情報保護法では1年以下の懲役または50万円以下の罰金という刑事罰があります。')
    script_line('さらに民事上の損害賠償、病院としての社会的信用の失墜——患者さんからの信頼を失ったら、この病院の存続にも関わります。')
    spacer(doc, 2)
    script_bold('★ 情報の保護は「ルールだから守る」ではなく、「患者さんとの信頼関係を守るために必要なこと」です。そう理解してほしいです。')

    # ===== ハラスメント予防 =====
    script_section('PART 02　ハラスメント予防（約8分）')
    script_note('【資料4ページ〜5ページ】')
    spacer(doc, 2)
    script_line('次にハラスメント予防です。まず種類を整理します。資料4ページをご覧ください。')
    spacer(doc, 2)
    script_line('パワハラは優越的な立場を使った言動。「なんでこんなこともできないんだ」という怒鳴り声は、指導ではなくパワハラです。', indent=True)
    script_line('セクハラは性的な言動。外見や体型についての不用意な発言も含まれます。', indent=True)
    script_line('マタハラは妊娠や育休に関わるもの。「このタイミングで妊娠するの？」は一言でアウトです。', indent=True)
    spacer(doc, 2)
    script_bold('★ そして今日メインでお話しするのが——カスハラです。')
    script_line('カスタマーハラスメント。顧客や患者、その家族から受けるハラスメントのことです。資料5ページを開いてください。')
    spacer(doc, 2)
    script_line('カスハラの定義を読みます。')
    script_bold('★「顧客・患者・その家族等が行う、社会通念上不相当な言動で、労働者の就業環境が害されるもの」。')
    script_line('これは2024年に厚生労働省が出した指針に記載された定義です。')
    spacer(doc, 2)
    script_line('病院は特にカスハラが起きやすい職場です。なぜなら患者さんやご家族は不安や不満を抱えた状態で来院します。その感情が時に、スタッフへの不当な言動に向かうことがあります。')
    spacer(doc, 2)
    script_line('どんな行為がカスハラに当たるか、具体例を見てください。')
    spacer(doc, 2)
    script_line('「訴えるぞ」「SNSで晒す」という脅し。大声での怒鳴り・罵倒。「上を呼べ」の繰り返し。これは暴言・脅迫系です。', indent=True)
    script_line('土下座や謝罪文の強要、医学的に不必要な処置の要求。これは不当要求系です。', indent=True)
    script_line('物を投げる、叩くという暴力行為。これは身体的行為。最悪の場合、刑事事件になります。', indent=True)
    script_line('そして長時間居座ってクレームを繰り返す。これも立派なカスハラです。', indent=True)
    spacer(doc, 2)
    script_bold('★ 「正当なクレーム」と「カスハラ」は区別してください。')
    script_line('患者さんからの不満や要望は、医療の質を改善する大事な声です。丁寧に対応しなければなりません。')
    script_line('でも、怒鳴る・脅す・暴力をふるうのは、正当なクレームではありません。「患者さんだから」と我慢し続ける必要はない。')
    spacer(doc, 2)
    script_line('では、カスハラに遭遇したらどうするか。対応フローを確認します。')
    spacer(doc, 2)
    script_line('まずステップ1。「これはクレームか、カスハラか」を判断する。怒鳴られているとパニックになりますが、一歩引いて状況を見てください。', indent=True)
    script_bold('★ ステップ2。一人で抱え込まない。すぐに同僚・先輩・リーダーへ声をかける。これが一番重要です。')
    script_line('ステップ3。複数人で対応する。必ず2名以上。これが大原則です。', indent=True)
    script_line('ステップ4。事実を記録する。日時・言動の内容を具体的にメモしておく。後で組織として対応するときに証拠になります。', indent=True)
    script_line('ステップ5。上長・相談窓口へ報告する。組織として対応する。', indent=True)
    spacer(doc, 2)
    script_line('2024年の法改正で、事業者はカスハラへの対策・相談体制の整備・被害スタッフへのフォローが義務とされました。')
    script_line('つまり、スタッフが一人で我慢するのは間違っています。組織で対応するのが正解です。病院としても、スタッフを守る体制を整備していきます。')
    spacer(doc, 2)
    script_bold('★ もう一度だけ言います。カスハラは「我慢するもの」ではありません。')

    # ===== 法令遵守 =====
    script_section('PART 03　法令遵守（約4分）')
    script_note('【資料6ページ】')
    spacer(doc, 2)
    script_line('最後に法令遵守、コンプライアンスについてです。')
    script_line('医療現場に関わる主な法令を資料でご確認ください。')
    spacer(doc, 2)
    script_line('個人情報保護法——これはPART01でお話しした通りです。漏洩時は72時間以内に当局への報告義務があります。', indent=True)
    script_line('医療法——医療記録の正確な保管義務と、インフォームドコンセントの確保が求められています。説明不足・記録の改ざんは医療法違反になります。', indent=True)
    script_line('労働施策総合推進法、いわゆるパワハラ防止法——これはハラスメントのパートで触れた通り、事業者に対策義務が課せられています。', indent=True)
    script_line('医師法・保健師助産師看護師法——業務上知った情報の守秘義務があります。退職後も続きます。先ほど言った話と同じですね。', indent=True)
    script_line('不正競争防止法——院内の業務秘密を持ち出す・漏らすことは禁じられています。', indent=True)
    spacer(doc, 2)
    script_line('違反するとどうなるか。刑事罰・行政処分・民事賠償の3つが待っています。')
    script_line('特に免許の停止・取り消しは、医療専門職としてのキャリアが終わることを意味します。')
    spacer(doc, 2)
    script_bold('★ そして一番伝えたいことは——「迷ったら報告してください」ということです。')
    script_line('問題が起きたとき、隠したくなる気持ちはわかります。でも隠すほど被害は広がります。早期報告が自分を・患者を・病院を守ります。')
    script_line('上長への相談が難しいときは、院内の相談窓口を使ってください。資料に窓口情報を記載してあります。')

    # ===== まとめ =====
    script_section('まとめ（約1分）')
    script_line('では最後にまとめます。資料の最終ページ、セルフチェックリストもご活用ください。')
    spacer(doc, 2)
    script_line('今日お伝えしたのは3つだけです。')
    spacer(doc, 2)
    script_bold('★ 1つ目。患者情報はあなたが「預かっている」もの。慎重に、目的の範囲内だけで扱ってください。')
    script_bold('★ 2つ目。ハラスメントは職場の問題。特にカスハラは一人で抱え込まず、必ず組織で対応してください。')
    script_bold('★ 3つ目。迷ったら報告。隠さないことが、自分と患者と病院を守ります。')
    spacer(doc, 2)
    script_line('研修後、セルフチェックリストを記入してみてください。チェックできない項目が、今日から変えるポイントです。')
    spacer(doc, 2)
    script_line('以上で研修を終わります。ご質問があればどうぞ。')

    # 時間メモ
    spacer(doc, 6)
    add_paragraph_border_bottom(doc.add_paragraph(''), sz=4)
    bordered_box(doc, '所要時間の目安',
        ['イントロ：約1分　／　個人情報保護：約6分',
         'ハラスメント予防：約8分　／　法令遵守：約4分',
         'まとめ：約1分',
         '合計：約20分（質疑応答で＋10分 = 合計約30分）'])

    # ナノバナナ指示
    page_break(doc)
    heading_para(doc, 'ナノバナナ　イラスト作成指示メモ', font_size=13)
    body_para(doc, '以下の4点をモノクロ・シンプルな線画スタイルで作成してください。\n完成したら 配布資料.docx の各プレースホルダー箇所に挿入してください。',
              font_size=10, color='444444')
    spacer(doc, 4)

    illust_items = [
        ('① 表紙用', '鍵・盾・法律の本を一つずつ持った3人の人物シルエット\n横並び・中央揃え・横150px×縦80px程度\nファイル名：cover_illust.png'),
        ('② 個人情報保護ページ用', '大きな鍵のかかったファイルボックスを両手で抱えている人物のシルエット\n右寄せ・横80px×縦80px程度\nファイル名：illust_privacy.png'),
        ('③ カスハラページ用', '左側に怒鳴っている人物（口大きく開け・腕を上げる）\n右側に盾を持って立つスタッフのシルエット\n二人の間に波線（衝撃・声を示す）\n横長・横200px×縦70px程度\nファイル名：illust_kushara.png'),
        ('④ 法令遵守ページ用', '天秤。左皿に病院の建物アイコン、右皿に法律書（本）のアイコン\nシンプルな線画・右寄せ・横90px×縦80px程度\nファイル名：illust_compliance.png'),
    ]
    for label, desc in illust_items:
        sub_heading(doc, label, font_size=10.5)
        for line in desc.split('\n'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            run.font.size = Pt(10)
            set_run_font(run)
        spacer(doc, 4)

    doc.save('C:/Users/yuuna/agens/勉強会/発表原稿.docx')
    print('発表原稿.docx を作成しました')


if __name__ == '__main__':
    create_handout()
    create_script()
    print('完了')
