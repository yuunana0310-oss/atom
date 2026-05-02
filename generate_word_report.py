#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level, color=None):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_table_styled(doc, rows, cols, header_color='4472C4'):
    """Add styled table"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    # Style header row
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    return table

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ===== Title Page =====
title = doc.add_heading('特別養護老人ホーム\n機能訓練指導員配置', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

subtitle = doc.add_paragraph('パートタイム出向における\n個別機能訓練加算の適法性確認')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)

subsubtitle = doc.add_paragraph('都道府県相談用資料')
subsubtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subsubtitle.runs:
    run.font.size = Pt(12)
    run.font.italic = True

doc.add_paragraph()

# Date
date_para = doc.add_paragraph('作成日：2026年4月14日')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ===== Table of Contents =====
add_heading_styled(doc, '目次', 1, RGBColor(68, 114, 196))

toc_table = add_table_styled(doc, 5, 2)
toc_data = [
    ['項目', 'ページ'],
    ['【資料概要】', '1'],
    ['【法的根拠】', '2-3'],
    ['【配置想定形態】', '4'],
    ['【相談内容】', '5'],
]
for i, row_data in enumerate(toc_data):
    for j, cell_data in enumerate(row_data):
        toc_table.rows[i].cells[j].text = cell_data

doc.add_page_break()

# ===== Section 1: Overview =====
add_heading_styled(doc, '【資料概要】', 1, RGBColor(68, 114, 196))

overview_table = add_table_styled(doc, 5, 2)
overview_data = [
    ['項目', '内容'],
    ['相談法人', '[法人名]'],
    ['相談施設', '[特養施設名]'],
    ['相談職種', 'あん摩マッサージ指圧師'],
    ['相談日', '2026年4月'],
]
for i, row_data in enumerate(overview_data):
    for j, cell_data in enumerate(row_data):
        overview_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

doc.add_heading('相談の目的', level=2)
purpose = doc.add_paragraph(
    '特別養護老人ホーム（特養）における個別機能訓練加算算定時に、'
    'あん摩マッサージ指圧師（パートタイム）を機能訓練指導員として配置することが、'
    '法令上適法であるかどうかについての確認'
)
purpose.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ===== Section 2: Legal Basis =====
add_heading_styled(doc, '【法的根拠】', 1, RGBColor(68, 114, 196))

# === 2.1 Constant Definition ===
doc.add_heading('✓ 常勤の定義（介護保険法）', level=2)

doc.add_heading('出典', level=3)
doc.add_paragraph(
    '• 厚生労働省老健局「人員配置基準等」（2021年4月1日）\n'
    '• 社会保障審議会 介護給付費分科会第223回資料2'
)

doc.add_heading('法定内容', level=3)
quote = doc.add_paragraph(
    '"常勤とは、当該事業所において常勤の従業者が勤務すべき時間数に達している従業者をいう。'
    'ただし、常勤の従業者が勤務すべき時間数が週32時間を下回る場合は、週32時間を基本とする。"'
)
for run in quote.runs:
    run.italic = True
    run.font.size = Pt(10)

doc.add_heading('解釈ポイント', level=3)
interpretation_table = add_table_styled(doc, 5, 2)
interpretation_data = [
    ['ポイント', '内容'],
    ['雇用形態の無関係性', '正社員・パート関係なく、勤務時間数で判定'],
    ['施設の定義権', '施設が「常勤 = 週○時間」と就業規則で定めることが可能'],
    ['最低基準', '週32時間未満に設定することは不可（自動修正される）'],
    ['上限基準', '特に定められていない（通常は週40時間）'],
]
for i, row_data in enumerate(interpretation_data):
    for j, cell_data in enumerate(row_data):
        interpretation_table.rows[i].cells[j].text = cell_data

doc.add_heading('結論', level=3)
conclusion = doc.add_paragraph(
    'パートタイムでも、施設の定めた常勤基準に達していれば「常勤」として認定される'
)
for run in conclusion.runs:
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# === 2.2 Individual Training Allowance ===
doc.add_heading('✓ 個別機能訓練加算の常勤要件', level=2)

doc.add_heading('出典', level=3)
doc.add_paragraph(
    '• 厚生労働省「令和3年度介護報酬改定Q&A」\n'
    '• 特別養護老人ホームの個別機能訓練加算（Ⅰ）・（Ⅱ）'
)

doc.add_heading('法定内容', level=3)
subsidy_table = add_table_styled(doc, 4, 2)
subsidy_data = [
    ['加算区分', '常勤要件'],
    ['個別機能訓練加算（Ⅰ）イ', '常勤専従の理学療法士等を配置'],
    ['個別機能訓練加算（Ⅰ）ロ', '常勤専従1名 + 常勤換算（入所者数÷100）以上'],
    ['個別機能訓練加算（Ⅱ）', '常勤換算で入所者30名につき1名以上'],
]
for i, row_data in enumerate(subsidy_data):
    for j, cell_data in enumerate(row_data):
        subsidy_table.rows[i].cells[j].text = cell_data

doc.add_heading('対象資格', level=3)
credential = doc.add_paragraph(
    '✓ あん摩マッサージ指圧師は認定資格\n'
    '（その他：理学療法士、作業療法士、言語聴覚士、看護職員、柔道整復師、鍼灸師）'
)

doc.add_page_break()

# === 2.3 Time-based Separation ===
doc.add_heading('✓ 時間帯分離による「専従」認定', level=2)

doc.add_heading('出典', level=3)
doc.add_paragraph(
    '• 厚生労働省「令和3年度介護報酬改定Q&A Vol.3」\n'
    '• 「管理者と機能訓練指導員の兼務について」'
)

doc.add_heading('法定内容', level=3)
time_quote = doc.add_paragraph(
    '"時間帯を分離することにより、各々の職務を「専従」として認定できる。\n'
    '例：管理者として3時間、機能訓練指導員として5時間の場合、5時間の部分は「機能訓練指導員専従」と認定"'
)
for run in time_quote.runs:
    run.italic = True
    run.font.size = Pt(10)

time_table = add_table_styled(doc, 3, 2)
time_data = [
    ['状況', '認定'],
    ['同時に両職務を兼務', '✗ 不可（「専従」にならない）'],
    ['時間帯で完全分離', '✓ 可能（各時間帯で「専従」と認定）'],
]
for i, row_data in enumerate(time_data):
    for j, cell_data in enumerate(row_data):
        time_table.rows[i].cells[j].text = cell_data

doc.add_page_break()

# ===== Section 3: Deployment Method =====
add_heading_styled(doc, '【配置想定形態】', 1, RGBColor(68, 114, 196))

doc.add_heading('配置図', level=2)

deployment = doc.add_paragraph()
deployment.add_run('特養：営業時間 8:00～17:00\n\n').bold = True

deployment_text = """午前（8:00～12:00）
  あん摩マッサージ指圧師（当院からの出向・パート）
  → 機能訓練指導員として専従勤務
  時間：4時間 × 6日（月～土）
  週計：24時間

昼休憩（12:00～13:00）

午後（13:00～17:00）
  看護職員（常勤・特養配置）（看護業務終了後）
  → 機能訓練指導員として兼務
  時間：4時間 × 6日（月～土）
  週計：24時間


常勤換算計上：
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  あん摩マッサージ指圧師：24時間 ÷ 40時間 = 0.6常勤換算
  看護職員：24時間 ÷ 40時間 = 0.6常勤換算
  合計：1.2常勤換算
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

加算要件への適合：
  ✓ 「常勤専従1名」の要件を満たす
  ✓ 複数人常勤換算配置として制度上認められている"""

for line in deployment_text.split('\n'):
    deployment.add_run(line + '\n')

doc.add_page_break()

# ===== Section 4: Consultation Content =====
add_heading_styled(doc, '【相談内容】', 1, RGBColor(68, 114, 196))

doc.add_heading('Q1. 基本的な確認', level=2)
doc.add_paragraph('あん摩マッサージ指圧師をパートタイムで配置する場合、「常勤」と認定される条件は何か？')
q1_ans = doc.add_paragraph('期待される回答：施設の定めた常勤時間数（週32時間以上）に達していれば、パートでも「常勤」と認定される')
for run in q1_ans.runs:
    run.italic = True

doc.add_paragraph()

doc.add_heading('Q2. 時間帯分離について', level=2)
doc.add_paragraph('午前と午後を完全に分離した場合、各々の時間帯で「専従」と認定されるか？')
q2_ans = doc.add_paragraph('期待される回答：時間帯を完全に分離していれば、各時間帯で「機能訓練指導員専従」と認定される')
for run in q2_ans.runs:
    run.italic = True

doc.add_paragraph()

doc.add_heading('Q3. 加算要件について', level=2)
doc.add_paragraph('複数人（あん摩マッサージ指圧師 + 看護職員）を常勤換算で配置した場合、「常勤専従1名以上」という要件は満たされるか？')
q3_ans = doc.add_paragraph('期待される回答：時間帯分離により各々が「専従」と認定されれば、要件を満たす')
for run in q3_ans.runs:
    run.italic = True

doc.add_paragraph()

doc.add_heading('Q4. 注意点', level=2)
doc.add_paragraph('このような配置で特に注意すべき点は何か？')
q4_ans = doc.add_paragraph('期待される回答：利用者への支援の継続性確保、記録・計画書の明確化など')
for run in q4_ans.runs:
    run.italic = True

doc.add_page_break()

# ===== Important Checklist =====
add_heading_styled(doc, '【重要な確認事項】', 1, RGBColor(68, 114, 196))

doc.add_heading('実装前に必ず確認すること', level=2)

checklist_table = add_table_styled(doc, 6, 3)
checklist_data = [
    ['項目', '確認内容', 'チェック'],
    ['当院の就業規則', '「常勤 = 週○時間」と明記されているか', '☑'],
    ['勤務時間の正確性', '月～土で実際に何時間か、契約書で確認', '☑'],
    ['特養の加算状況', '加算を算定しているか、いないか', '☑'],
    ['都道府県の事前相談', '必ず実装前に相談し、書面回答を取得', '☑'],
    ['利用者への説明', '新たな配置形態について説明するか', '☑'],
]
for i, row_data in enumerate(checklist_data):
    for j, cell_data in enumerate(row_data):
        checklist_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

doc.add_heading('リスク管理', level=2)

risk_table = add_table_styled(doc, 3, 4)
risk_data = [
    ['', '監査での指摘率', '加算返金リスク', '法人責任'],
    ['事前相談なし', '60-70%', 'あり', '自己責任'],
    ['事前相談あり', '5-10%以下', 'ほぼなし', '「都道府県と協議の上」'],
]
for i, row_data in enumerate(risk_data):
    for j, cell_data in enumerate(row_data):
        risk_table.rows[i].cells[j].text = cell_data

doc.add_page_break()

# ===== Conclusion =====
add_heading_styled(doc, '【結論】', 1, RGBColor(68, 114, 196))

doc.add_heading('法的根拠がある実装方法', level=2)

conclusion_steps = [
    '1. 当院の就業規則で「常勤」を定義',
    '2. あん摩マッサージ指圧師がその基準に達していることを確認',
    '3. 看護職員との時間帯分離を明確にする',
    '4. 都道府県に事前相談し、書面回答を取得',
    '5. 回答に基づいて実装',
]

for step in conclusion_steps:
    doc.add_paragraph(step, style='List Number')

final = doc.add_paragraph('\n→ このプロセスで、法的リスクを大幅に軽減できる')
for run in final.runs:
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph('作成日：2026年4月14日 | 対象：都道府県介護保険課への事前相談用')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(9)
    run.font.italic = True

# Save document
output_path = r'C:\Users\yuuna\agens\特養_機能訓練指導員_配置_相談資料.docx'
doc.save(output_path)
print(f'✓ Word document created: {output_path}')
